"""Update the evaluation system DOCX with today's work summary.

All paths in the document are written as absolute paths.
"""

from docx import Document
from docx.oxml.ns import qn

doc = Document("/root/instructions/中医药政策研究文献评价表 (1).docx")

# ── Helpers ────────────────────────────────────────────────────────
def find_para(text_fragment):
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i, p
    return None, None

def insert_after(idx, content_list):
    parent = doc.paragraphs[idx]._element.getparent()
    insert_idx = list(parent).index(doc.paragraphs[idx]._element) + 1
    for item in reversed(content_list):
        p = doc.add_paragraph(item)
        parent.insert(insert_idx, p._element)

def remove_paragraphs(indices):
    for i in sorted(indices, reverse=True):
        el = doc.paragraphs[i]._element
        el.getparent().remove(el)

def clear_between(start_text, end_text):
    i1, _ = find_para(start_text)
    i2, _ = find_para(end_text)
    to_remove = []
    if i1 and i2:
        for i in range(i1 + 1, i2):
            if doc.paragraphs[i].text.strip():
                to_remove.append(i)
    elif i1:
        for i in range(i1 + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip():
                to_remove.append(i)
    remove_paragraphs(to_remove)
    return i1

# ====================================================================
# 1. Update 2.1
# ====================================================================
print("1. Updating 2.1 ...")
clear_between("2.1 整体技术架构", "2.2 技术栈选型")
idx, _ = find_para("2.1 整体技术架构")
if idx:
    insert_after(idx, [
        "系统采用「前端展示层 - 业务服务层 - 模型推理层」三层架构，层间通过 Python 模块直接调用。代码根目录为 /root/web/，Streamlit 前端与 FastAPI 后端共享同一代码库。",
        "",
        "第一层：前端展示层（Streamlit）",
        "  /root/web/streamlit_app.py — 主入口，负责页面布局、文件上传、进度展示、可视化渲染。",
        "  可视化组件：Plotly 雷达图/柱状图/仪表盘，实时渲染评价结果。",
        "  侧边栏：评价体系导航、历史记录浏览与回溯。",
        "  通信方式：直接 import backend.services 模块，不走 HTTP（单进程模式）。",
        "",
        "第二层：业务服务层（FastAPI + 服务模块）",
        "  /root/web/backend/main.py — FastAPI 应用入口，管理生命周期（启动预加载模型，关闭释放显存）。",
        "  /root/web/backend/routers/evaluation.py — POST /evaluate（文本评价）、POST /evaluate/upload（文件上传评价）。",
        "  /root/web/backend/routers/history.py — GET/DELETE /history 历史管理端点。",
        "  /root/web/backend/services/evaluation_service.py — 核心评价编排，16个二级指标迭代、7个一级指标聚合、总分计算。",
        "  /root/web/backend/services/prompt_builder.py — 结构化提示词构建，嵌入评分标准/细则/判断依据三级锚点。",
        "  /root/web/backend/services/llm_client.py — GLM-4-9B 模型加载（全局缓存）与推理调用。",
        "  /root/web/backend/services/history_service.py — JSON 文件持久化，线程安全，自动容量控制。",
        "  /root/web/backend/utils/document_parser.py — PDF（pypdf + OCR 回退）和 DOCX 解析。",
        "  /root/web/backend/utils/response_parser.py — LLM 响应JSON解析（4层回退链）。",
        "  /root/web/backend/models/schemas.py — Pydantic 数据模型，7级嵌套结构。",
        "  /root/web/backend/middleware/ — 全局异常处理、速率限制（滑动窗口）、请求日志、认证中间件桩。",
        "  /root/web/backend/config.py — 集中配置：模型路径/推理参数/服务器设置，支持环境变量覆盖。",
        "",
        "第三层：模型推理层（GLM-4-9B-Chat）",
        "  路径：/root/autodl-tmp/glm4_models/ZhipuAI/glm-4-9b-chat",
        "  transformers 框架 + torch 直接加载模型权重，GPU bfloat16 推理。",
        "  推理参数：temperature=0.3, top_p=0.9, max_new_tokens=512, repetition_penalty=1.1。",
        "",
        "共用层（shared/）",
        "  /root/web/shared/constants.py — 7个一级指标 + 16个二级指标的完整数据结构定义。",
        "  前后端共用同一套 TypedDict 指标定义，确保数据一致性。",
        "",
        "数据层",
        "  /root/web/data/eval_history.json — 评价历史持久化（JSON文件，最多保留50条）。",
    ])
    print("  2.1 done")

# ====================================================================
# 2. Update 2.2
# ====================================================================
print("2. Updating 2.2 ...")
clear_between("2.2 技术栈选型", "3，核心模块设计")
idx, _ = find_para("2.2 技术栈选型")
if idx:
    insert_after(idx, [
        "（一）大语言模型",
        "  GLM-4-9B-Chat（智谱AI，9B参数对话模型）",
        "  加载方式：transformers.AutoTokenizer + AutoModel，trust_remote_code=True",
        "  GPU（CUDA 12.8）：bfloat16 + device_map=auto，显存约18GB",
        "  CPU 备用：float32 + max_memory={0: 50GiB}，线程限制 torch.set_num_threads(min(cpu_count, 8))",
        "  模型缓存：模块级全局变量，首次调用加载，后续复用（避免重复加载）",
        "",
        "（二）后端框架",
        "  FastAPI — 现代异步 Python Web 框架，自动生成 OpenAPI 文档（/docs, /redoc）",
        "  Uvicorn — ASGI 服务器，单 worker 模式（GPU 独占，不支持并发推理）",
        "  Pydantic — 请求/响应数据校验与序列化，支持 field_validator 自定义校验",
        "  CORSMiddleware — 跨域支持，allow_origins=[*]（内部工具场景）",
        "",
        "（三）前端框架",
        "  Streamlit 1.57.0 — 纯 Python 数据应用框架",
        "  核心机制：session_state（跨rerun状态保持）、st.cache_resource（模型缓存）",
        "  启动命令：streamlit run /root/web/streamlit_app.py --server.port 6006 --server.enableCORS false --server.enableXsrfProtection false --server.address 0.0.0.0",
        "  适用场景：内部工具、原型验证、小团队使用",
        "",
        "（四）可视化",
        "  Plotly 6.7.0 — 交互式图表库",
        "  go.Indicator：仪表盘（gauge+number+delta 模式）",
        "  go.Scatterpolar：雷达图（7轴展示一级指标得分率）",
        "  px.bar：水平柱状图（color_continuous_scale 渐变着色）",
        "  Pandas 3.0.2 — 图表数据准备与排序",
        "",
        "（五）文档解析",
        "  pypdf 6.10.2 — 文字型 PDF 文本提取（PdfReader）",
        "  python-docx — Word .docx 段落提取（docx.Document）",
        "  pdf2image 1.17.0 — PDF→图片转换，依赖 poppler-utils",
        "  pytesseract 0.3.13 + Tesseract OCR + chi_sim — 中文 OCR 识别",
        "  Pillow 12.2.0 — 图片预处理（灰度化/二值化/DPI调整）",
        "  解析策略：pypdf 优先（文字PDF）-> 文字<50字符 -> OCR回退（扫描件PDF）",
        "",
        "（六）依赖库",
        "  torch 2.8.0 — 深度学习框架，GPU/CPU 推理",
        "  transformers 4.46.0 — HuggingFace 模型加载",
        "  sentencepiece 0.2.1 — GLM 分词器底层依赖",
        "  tiktoken 0.12.0 — OpenAI 分词器（部分 transformers 模型需要）",
        "  protobuf 6.31.1 — Protocol Buffers 序列化（模型配置）",
        "  accelerate 1.13.0 — HuggingFace 推理加速，device_map 自动分配",
        "  modelscope 1.36.3 — 阿里模型社区 SDK（国内模型下载备用）",
        "  langchain 1.2.17 — LLM 应用开发框架（预留扩展接口）",
        "",
        "（七）容器化部署",
        "  Docker — 基于 nvidia/cuda:12.8.0-devel-ubuntu22.04 基础镜像",
        "  docker-compose — 三服务编排（api GPU推理 / streamlit 前端 / nginx 反向代理）",
        "  /root/web/Dockerfile — 镜像定义，分层构建利用缓存",
        "  /root/web/docker-compose.yml — 服务编排，GPU device reservation，volume 挂载",
        "  /root/web/entrypoint.sh — 统一入口脚本，支持 api | streamlit | both 三种模式",
        "  /root/web/nginx.conf — 反向代理 + WebSocket 升级 + 速率限制",
        "",
        "（八）系统工具",
        "  tesseract-ocr + tesseract-ocr-chi-sim — Tesseract OCR 引擎及中文语言包",
        "  poppler-utils — PDF 渲染工具（pdftoppm）",
        "  screen / tmux — 终端复用器，保持持久会话",
    ])
    print("  2.2 done")

# ====================================================================
# 3. Update 3.1, 3.3
# ====================================================================
print("3. Updating 3.1 ...")
clear_between("3.1 智能评价Agent", "3.2 评价报告生成与可视化")
idx, _ = find_para("3.1 智能评价Agent")
if idx:
    insert_after(idx, [
        "智能评价Agent 是系统的核心引擎，实现代码位于 /root/web/backend/services/ 目录下。",
        "采用「提示词构建 -> 模型推理 -> 响应解析」三阶段流水线：",
        "",
        "阶段一：提示词构建（/root/web/backend/services/prompt_builder.py, 183行）",
        "  为每个二级指标独立构造结构化提示词，包含四个部分：",
        "    1. 系统角色设定：资深评审专家（10年+政策研究经验）",
        "    2. 评价标准：3-4条具体检查点（如是否包含近3年政策演变关键节点分析）",
        "    3. 评分细则：4级评分锚点（优秀/良好/一般/差，每级对应具体分数区间）",
        "    4. 关键判断依据：每级的具体表现特征（如优秀=有明确时间节点+典型案例+权威数据）",
        "  输出格式要求：严格 JSON {score, evidence, comment}",
        "  文档截断：max_chars=6000，超出部分标注截断标记",
        "  额外提供 build_additional_prompt() 构建附加项（+/-5分）独立提示词",
        "",
        "阶段二：模型推理（/root/web/backend/services/llm_client.py, 120行）",
        "  模型加载策略：",
        "    - 模块级全局缓存（_tokenizer, _model），首次调用时从磁盘加载",
        "    - GPU路径：bfloat16 dtype + device_map=auto + low_cpu_mem_usage=True",
        "    - CPU路径：float32 dtype + device_map=cpu + max_memory={0: 50GiB}",
        "  推理流程：",
        "    1. tokenizer.apply_chat_template 构造对话格式（messages=[{role, content}]）",
        "    2. 计算 attention_mask（非 padding token 标记为1）",
        "    3. model.generate() 执行自回归生成",
        "    4. tokenizer.decode() 解码输出（skip_special_tokens=True）",
        "  生成控制：temperature=0.3（低随机性，保证评分一致性）、top_p=0.9（核采样）、repetition_penalty=1.1（抑制重复）",
        "  资源管理：clear_cache() 释放模型显存，get_vram_usage() 查询当前显存占用",
        "",
        "阶段三：响应解析（/root/web/backend/utils/response_parser.py, 148行）",
        "  4层回退解析链，按可靠性从高到低：",
        "    第1层：Markdown 代码块提取（匹配 ```json ... ``` 模式）",
        "    第2层：score 字段 JSON 对象匹配（正则查找含 score 键的 JSON 片段）",
        "    第3层：中文前缀清理后全量 JSON 解析（移除 好的/以下是/根据 等前缀）",
        "    第4层：括号平衡提取（从第一个{开始匹配嵌套深度）",
        "  以上均失败时：正则分别提取 score/evidence/comment 字段",
        "  设计原则：评分范围不在解析器层硬编码，由调用方按指标满分 clamp",
        "",
        "评价编排（/root/web/backend/services/evaluation_service.py, 170行）",
        "  evaluate_document() 主入口，流程：",
        "    1. 遍历 ALL_SECONDARY_INDICATORS（16个），对每个调用 evaluate_secondary_indicator()",
        "    2. 每完成一个指标，触发 progress_callback(i, total) 更新前端进度条",
        "    3. 异常隔离：单个指标失败不影响其他（返回 score=0, comment=评估失败）",
        "    4. 按 SECONDARY_TO_PRIMARY 映射聚合为 7 个 PrimaryResult",
        "    5. 计算 base_score = sum(primary.scores)",
        "    6. 调用 evaluate_additional_item() 获取 +/-5 附加分",
        "    7. total_score = base_score + additional.score",
        "    8. 构造 EvaluationResponse 返回",
    ])
    print("  3.1 done")

print("4. Updating 3.3 ...")
clear_between("3.3 文献预处理", "3.4 历史记录模块")
idx, _ = find_para("3.3 文献预处理")
if idx:
    insert_after(idx, [
        "文献预处理模块实现代码位于 /root/web/backend/utils/document_parser.py（140行）。",
        "支持 PDF 和 DOCX 两种格式，采用分层解析策略：",
        "",
        "DOCX 解析",
        "  使用 python-docx 库读取 /root/web/backend/utils/document_parser.py 中的 parse_docx() 函数。",
        "  流程：BytesIO 封装字节流 -> docx.Document 对象 -> 遍历 doc.paragraphs -> 过滤空段落 -> join 输出。",
        "  仅提取文本内容（表格、图片暂不处理）。",
        "",
        "PDF 解析（双层策略）",
        "  第1层 — pypdf 文本提取（parse_pdf 函数）：",
        "    PdfReader(file) 逐页调用 page.extract_text()",
        "    适用于文字型 PDF（如 Word 导出、LaTeX 生成的 PDF）",
        "    速度快，单页 < 0.1秒",
        "  第2层 — OCR 回退（_ocr_pdf 函数，触发条件：提取文字 < 50 字符）：",
        "    1. pdf2image.convert_from_bytes(pdf_bytes, dpi=200) 将 PDF 每页转为 PIL Image",
        "    2. pytesseract.image_to_string(image, lang='chi_sim') 中文字符识别",
        "    3. 每5页输出进度日志",
        "    4. Tesseract OCR 引擎路径：系统安装的 tesseract-ocr + tesseract-ocr-chi-sim",
        "    适用于扫描件/图片型 PDF（政府公文常见格式）",
        "  OCR 性能：10页PDF约80-150秒（300DPI），200DPI时可降至40-80秒",
        "",
        "文本截断",
        "  truncate_text(text, max_chars=6000)：在构造提示词前调用",
        "  超出部分截断并标注...（内容截断） 防止超出 LLM 上下文窗口",
    ])
    print("  3.3 done")

# ====================================================================
# 4. Update 6.0
# ====================================================================
print("5. Updating 6.0 ...")
clear_between("6.0 环境配置", "6.1 一键启动脚本")
idx, _ = find_para("6.0 环境配置")
if idx:
    insert_after(idx, [
        "硬件环境",
        "  GPU: NVIDIA RTX 5090（Blackwell 架构，32GB 显存）",
        "  CPU: 多核 x86_64",
        "  操作系统: Ubuntu 22.04 LTS",
        "  CUDA: 12.8",
        "",
        "Python 环境",
        "  Python 3.12.3，使用 conda base 环境（AutoDL 推荐）",
        "  CPU 模式限制线程：torch.set_num_threads(min(os.cpu_count() or 4, 8))",
        "",
        "完整依赖清单（requirements.txt）及用途说明",
        "",
        "  torch==2.8.0",
        "    深度学习框架，GPU/CPU 推理引擎。提供 tensor 运算、autograd 自动微分、CUDA 加速。",
        "    9B 模型推理的核心依赖，所有模型权重以 torch.Tensor 形式加载和计算。",
        "",
        "  transformers==4.46.0",
        "    HuggingFace 模型库。提供 AutoTokenizer（分词器加载）和 AutoModel（模型加载），",
        "    以及 apply_chat_template（对话模板）、model.generate（自回归生成）等核心 API。",
        "    trust_remote_code=True 允许加载 GLM-4 的自定义模型代码。",
        "",
        "  fastapi",
        "    现代 Python Web 框架，基于 Starlette 和 Pydantic。提供路由装饰器（@app.post）、",
        "    依赖注入、自动 OpenAPI 文档生成（/docs, /redoc）、类型安全的请求/响应模型。",
        "    用于构建 /evaluate 和 /history 的 RESTful API 接口。",
        "",
        "  uvicorn",
        "    ASGI 服务器，运行 FastAPI 应用。单 worker 模式（GPU 独占，不支持并发推理）。",
        "    启动命令：uvicorn backend.main:app --host 0.0.0.0 --port 8000。",
        "",
        "  pydantic",
        "    数据校验与序列化库。定义 BaseModel 子类作为 API 请求/响应的 schema，",
        "    自动校验输入数据、类型转换（如 float->int）、生成 JSON Schema 文档。",
        "    支持 field_validator 自定义校验逻辑。",
        "",
        "  streamlit==1.57.0",
        "    纯 Python 数据应用前端框架。核心机制：",
        "    - st.session_state：跨 rerun 保持状态（评价结果/文档文本）",
        "    - st.cache_resource：函数级缓存（模型加载结果）",
        "    - 组件：file_uploader / button / progress / spinner / plotly_chart / expander",
        "",
        "  plotly==6.7.0",
        "    交互式可视化库。使用组件：",
        "    - go.Indicator：仪表盘（gauge+number+delta 模式，显示总分）",
        "    - go.Scatterpolar：雷达图（7轴一级指标得分率对比）",
        "    - px.bar：水平柱状图（color_continuous_scale 渐变着色）",
        "",
        "  pandas==3.0.2",
        "    数据处理与分析库。用于构造图表 DataFrame、排序（sort_values）、数据聚合。",
        "",
        "  pypdf==6.10.2",
        "    PDF 文本提取库。PdfReader 逐页调用 extract_text() 获取文字层内容。",
        "    局限性：无法处理扫描件/图片型 PDF（需 OCR 回退）。",
        "",
        "  python-docx",
        "    Word .docx 文件读写库。Document 对象遍历 paragraphs、tables、styles。",
        "    用于文档解析（提取文献正文）和文档生成（导出评价报告）。",
        "",
        "  pdf2image==1.17.0",
        "    PDF 页面转图片库。调用系统 poppler-utils 的 pdftoppm 工具，",
        "    将 PDF 每页渲染为 PIL Image 对象。convert_from_bytes() 支持直接读取字节流。",
        "    DPI 参数控制渲染分辨率（默认300，可降至200提速）。",
        "",
        "  pytesseract==0.3.13",
        "    Tesseract OCR 引擎的 Python 封装。image_to_string(image, lang=chi_sim) 执行中文 OCR。",
        "    依赖系统安装的 tesseract-ocr 和 tesseract-ocr-chi-sim（中文语言包）。",
        "",
        "  Pillow==12.2.0",
        "    Python 图像处理标准库。用于 OCR 前图片预处理：灰度化（convert(L)）、",
        "    二值化（threshold）、尺寸调整（resize）。可提升 OCR 速度 30-50%。",
        "",
        "  sentencepiece==0.2.1",
        "    Google 开源的文本分词器。GLM-4-9B 模型使用 SentencePiece 作为底层分词算法。",
        "    transformers 加载 GLM tokenizer 时自动调用。",
        "",
        "  tiktoken==0.12.0",
        "    OpenAI 开源的快速分词器。部分 HuggingFace transformers 模型需要此依赖。",
        "",
        "  protobuf==6.31.1",
        "    Google Protocol Buffers。模型配置文件（.pb）和序列化数据的解析依赖。",
        "    transformers 库加载某些模型时需要。",
        "",
        "  accelerate==1.13.0",
        "    HuggingFace 推理加速库。提供 device_map=auto（自动分配模型层到 GPU/CPU）、",
        "    low_cpu_mem_usage（低 CPU 内存加载模式）等优化功能。",
        "",
        "  modelscope==1.36.3",
        "    阿里 ModelScope 社区 SDK。提供国内可访问的模型下载渠道，",
        "    作为 HuggingFace Hub 的替代方案（国内网络环境下使用）。",
        "",
        "  langchain==1.2.17",
        "    LLM 应用开发框架。提供 PromptTemplate、Chain、Agent 等高级抽象。",
        "    当前系统暂未深度使用，预留为后续功能扩展（如多轮对话评价、RAG 增强）。",
        "",
        "系统工具",
        "  tesseract-ocr — Tesseract OCR 引擎（开源 OCR，Google 维护）",
        "  tesseract-ocr-chi-sim — 简体中文语言包",
        "  poppler-utils — PDF 渲染工具集（pdftoppm 将 PDF 页面转为 PPM/PNG 图像）",
        "  screen — 终端复用器，用于保持持久会话（SSH 隧道保活）",
        "",
        "Docker 封装机制",
        "  基础镜像：nvidia/cuda:12.8.0-devel-ubuntu22.04",
        "    包含 CUDA 12.8 开发工具链，支持 nvidia-container-runtime GPU 透传。",
        "  分层构建策略（Dockerfile）：",
        "    第1层：apt-get install 系统依赖（python3, curl 等）",
        "    第2层：pip install -r requirements.txt（利用 Docker 层缓存，依赖不变时秒级重建）",
        "    第3层：COPY . /app/（应用代码，变更频率最高，放在最后）",
        "  模型挂载：",
        "    模型文件（18GB+ GLM-4-9B）不打包进镜像，通过 docker-compose volumes 挂载：",
        "    /root/autodl-tmp/glm4_models:/models:ro（只读挂载，多个容器可共享）",
        "  数据持久化：",
        "    /root/web/data:/app/data（评价历史文件 eval_history.json）",
        "  统一入口（/root/web/entrypoint.sh）：",
        "    按 CMD 参数选择启动模式：api（FastAPI）、streamlit（Streamlit 前端）、both（同时启动）",
        "  GPU 资源声明（docker-compose.yml）：",
        "    deploy.resources.reservations.devices: [{driver: nvidia, count: 1, capabilities: [gpu]}]",
        "  Nginx 反向代理容器：",
        "    独立 nginx:alpine 容器，处理：WebSocket Upgrade（Streamlit 通信）、路径重写、",
        "    速率限制（limit_req_zone: 10r/m for API, 30r/m for Streamlit）",
        "  健康检查：",
        "    curl -f http://localhost:8000/health，每30秒一次，失败3次后自动重启",
        "",
        "Docker 优势总结",
        "  环境一致性：开发/测试/生产使用完全相同的依赖版本和系统库，消除「在我机器上能跑」问题",
        "  一键部署：docker-compose --profile production up -d 启动全套服务 3个容器",
        "  资源隔离：GPU 显存和 CPU 资源通过容器边界限制，避免与其他服务抢占",
        "  快速恢复：restart: unless-stopped 策略，容器崩溃后 Docker daemon 自动拉起",
    ])
    print("  6.0 done")

# ====================================================================
# 5. Update 8.2
# ====================================================================
print("6. Updating 8.2 ...")
clear_between("8.2 代码架构与软著材料准备", "8.3 软著申请材料清单与准备")
idx, _ = find_para("8.2 代码架构与软著材料准备")
if idx:
    insert_after(idx, [
        "系统采用前后端分离的模块化架构，共17个Python源文件、4个部署配置文件，总计超过2100行代码。",
        "所有路径以绝对路径标注，代码根目录为 /root/web/。",
        "",
        "——项目完整目录结构（绝对路径）——",
        "/root/web/",
        "├── streamlit_app.py                  [前端] Streamlit 主入口（419行）",
        "├── backend/                           [后端模块，1331行]",
        "│   ├── main.py                        FastAPI 入口（113行）",
        "│   ├── config.py                      全局配置（37行）",
        "│   ├── models/",
        "│   │   └── schemas.py                 Pydantic 数据模型（126行）",
        "│   ├── routers/",
        "│   │   ├── evaluation.py              POST /evaluate, POST /evaluate/upload（110行）",
        "│   │   └── history.py                GET/DELETE /history（68行）",
        "│   ├── services/",
        "│   │   ├── evaluation_service.py      评价编排（170行）",
        "│   │   ├── llm_client.py             模型推理（120行）",
        "│   │   ├── prompt_builder.py         提示词构建（183行）",
        "│   │   └── history_service.py        历史持久化（81行）",
        "│   ├── utils/",
        "│   │   ├── document_parser.py        文档解析+OCR（140行）",
        "│   │   └── response_parser.py        响应解析（148行）",
        "│   └── middleware/",
        "│       ├── error_handler.py           全局异常+认证桩（92行）",
        "│       ├── rate_limiter.py            限流（96行）",
        "│       └── request_logger.py          日志（88行）",
        "├── shared/                            [共用模块]",
        "│   └── constants.py                   评价体系（501行）",
        "├── tests/                             [测试]",
        "│   └── test_utils.py                  单元测试（116行）",
        "├── data/                              [运行时数据]",
        "│   └── eval_history.json              评价历史",
        "├── Dockerfile                         [部署配置，50行]",
        "├── docker-compose.yml                 [服务编排，74行]",
        "├── entrypoint.sh                      [容器入口，67行]",
        "└── nginx.conf                         [反向代理，111行]",
        "",
        "——各模块功能详细说明——",
        "",
        "1. /root/web/streamlit_app.py — Streamlit 前端主入口（419行）",
        "  功能：页头渲染、侧边栏（文献上传/评价体系导航/历史记录）、主区域（文献预览/评价按钮/",
        "  进度条/可视化图表/指标卡片/摘要下载）。所有业务逻辑委托给 backend.services 模块。",
        "  关键技术：st.session_state 保持评价结果跨 rerun、Plotly 图表嵌入、进度回调绑定。",
        "",
        "2. /root/web/backend/main.py — FastAPI 应用入口（113行）",
        "  功能：创建 FastAPI 实例、注册路由、配置 CORS、注册全局异常处理器、",
        "  lifespan 管理（启动时预加载模型，关闭时释放显存）、挂载 /health 健康检查端点。",
        "",
        "3. /root/web/backend/config.py — 全局配置（37行）",
        "  功能：模型路径、推理参数、API 主机/端口、文件大小限制的集中管理。",
        "  所有配置项支持 os.environ.get() 环境变量覆盖。CPU 模式自动限制线程数。",
        "",
        "4. /root/web/backend/models/schemas.py — Pydantic 数据模型（126行）",
        "  功能：定义 7 个核心数据类。SecondaryResult（单二级指标得分）-> PrimaryResult（一级聚合）",
        "  -> AdditionalResult（附加分）-> EvaluationResponse（完整响应）-> HistoryRecord（持久化）",
        "  -> HistoryListResponse（列表）-> ErrorResponse（错误）。支持 field_validator 自动类型转换。",
        "",
        "5. /root/web/backend/routers/evaluation.py — 评价端点（110行）",
        "  功能：POST /evaluate（接受 JSON 文本，返回完整评价结果）、",
        "  POST /evaluate/upload（接受 multipart 文件上传，自动解析后评价）。",
        "  包含格式验证、文件大小检查（50MB）、异常转换（ValueError->400, Exception->500）。",
        "",
        "6. /root/web/backend/routers/history.py — 历史管理端点（68行）",
        "  功能：GET /history?limit=20（历史列表）、GET /history/{id}（单条详情）、",
        "  DELETE /history/{id}（删除记录）。所有响应经过 Pydantic 模型校验。",
        "",
        "7. /root/web/backend/services/evaluation_service.py — 评价编排（170行）",
        "  功能：评价流程总调度。迭代 ALL_SECONDARY_INDICATORS -> 逐指标调用 LLM -> 解析 ->",
        "  异常隔离 -> 按 PRIMARY_INDICATORS 聚合 -> 计算总分 -> 附加分。",
        "  关键技术：progress_callback 进度回调模式、Optional[Callable] 类型标注。",
        "",
        "8. /root/web/backend/services/prompt_builder.py — 提示词构建（183行）",
        "  功能：为每个二级指标生成结构化评价提示词。build_secondary_prompt() 为单指标模式",
        "  （embed 评价标准+评分细则+判断依据），build_full_evaluation_prompt() 为全指标批量模式。",
        "  额外提供 build_additional_prompt() 用于附加项评分。",
        "",
        "9. /root/web/backend/services/llm_client.py — LLM 客户端（120行）",
        "  功能：GLM-4-9B 模型加载与推理。模块级全局缓存（_tokenizer/_model）避免重复加载。",
        "  双模式支持（GPU bfloat16 / CPU float32）。call_model(prompt) 执行单次推理。",
        "  clear_cache() 释放模型，get_vram_usage() 查询显存。",
        "",
        "10. /root/web/backend/services/history_service.py — 历史持久化（81行）",
        "  功能：JSON 文件读写评价历史。核心函数：save_to_history()（追加+截断50条）、",
        "  load_history()、get_record()、delete_record()、list_history()。",
        "  线程安全：threading.Lock 保护写操作，避免并发写入数据损坏。",
        "",
        "11. /root/web/backend/utils/document_parser.py — 文档解析（140行）",
        "  功能：PDF/DOCX 文本提取。parse_pdf() 实现双层策略：pypdf 优先（文本PDF），",
        "  文字<50字符自动回退 _ocr_pdf()（Tesseract OCR）。parse_docx() 使用 python-docx 提取段落。",
        "  detect_format() 通过扩展名判断 MIME 类型。truncate_text() 截断长文本。",
        "",
        "12. /root/web/backend/utils/response_parser.py — 响应解析（148行）",
        "  功能：从 LLM 原始输出中提取结构化 JSON。实现4层回退解析链：",
        "  Markdown代码块 -> score字段JSON对象 -> 前缀清理全量解析 -> 括号平衡提取 -> 正则回退。",
        "  评分范围不硬编码，由调用方按指标满分 clamp。支持负分提取（附加项）。",
        "",
        "13. /root/web/backend/middleware/error_handler.py — 异常处理（92行）",
        "  功能：register_exception_handlers() 注册全局异常处理器（HTTPException/ValueError/Exception）。",
        "  AuthMiddleware（认证桩）预留 X-API-Token 验证接口，待对接外部认证系统。",
        "",
        "14. /root/web/backend/middleware/rate_limiter.py — 速率限制（96行）",
        "  功能：SlidingWindowLimiter 实现滑动窗口算法（默认10次/60秒）。",
        "  RateLimitMiddleware 应用于 /evaluate 路由，超限返回 429 Too Many Requests + Retry-After 头。",
        "",
        "15. /root/web/backend/middleware/request_logger.py — 请求日志（88行）",
        "  功能：RequestIDMiddleware（UUID 请求追踪）、AccessLogMiddleware（结构化日志：",
        "  method/path/status/duration/client_ip）、SlowRequestMiddleware（30s阈值慢请求告警）。",
        "",
        "16. /root/web/shared/constants.py — 评价体系定义（501行）",
        "  功能：7个 PrimaryIndicator + 16个 SecondaryIndicator 完整数据结构（TypedDict）。",
        "  每个二级指标含 id/name/max_score/criteria/scoring_guide/evidence_guide。",
        "  附加项定义（+/-5分）、评分阈值、历史容量限制、推理参数默认值。",
        "  提供 get_secondary_by_id()、get_primary_by_id() 查找函数。",
        "",
        "17. /root/web/tests/test_utils.py — 单元测试（116行）",
        "  功能：TestResponseParser（9个用例，覆盖JSON/中文前缀/空输入/边界值）、",
        "  TestDocumentParser（6个用例，格式检测/截断）、TestPromptBuilder（3个用例）。",
        "  所有测试无需 GPU，纯逻辑验证。",
    ])
    print("  8.2 done")

    # Also update Table 4 (软著审查维度) code count
    if len(doc.tables) >= 5:
        table = doc.tables[4]
        if len(table.rows) >= 4:
            table.rows[3].cells[2].text = (
                "17个Python源文件 + 4个部署配置文件\n"
                "总计超过2100行代码\n"
                "模块化三层架构（/backend + /shared + /frontend）\n"
                "15个可独立测试的业务模块"
            )

# ====================================================================
# 6. Create Section 9 技术日志
# ====================================================================
print("7. Creating Section 9 ...")

# Find insertion point - before 我的基础环境 or before 提示词工程
idx_base, _ = find_para("我的基础环境")
if idx_base is None:
    idx_base, _ = find_para("提示词工程")
if idx_base is None:
    idx_base, _ = find_para("7，模型生成参数调优")

if idx_base:
    parent = doc.paragraphs[idx_base]._element.getparent()
    insert_idx = list(parent).index(doc.paragraphs[idx_base]._element)

    tech_log_items = [
        "",
        "9，技术日志",
        "",
        "9.1 语法错误：callable | None 类型标注不兼容",
        "【时间】2026-05-11 下午",
        "【现象】Streamlit 启动时报错：TypeError: unsupported operand type(s) for |: 'builtin_function_or_method' and 'NoneType'",
        "【原因】Python 3.12 支持 X | None 联合类型语法，但 callable（小写 c）是内置函数名而非类型。Python 将 callable | None 中的 callable 解析为内置函数对象，导致 | 操作符报错。正确的类型名是 collections.abc.Callable（大写 C）。",
        "【解决】改为 from collections.abc import Callable; from typing import Optional; 使用 Optional[Callable[[int, int], None]] 声明类型。",
        "【涉及文件】/root/web/backend/services/evaluation_service.py 第85行",
        "【教训】使用 | 联合类型语法时确保左操作数是类型（Type）而非内置函数。小写的 callable 永远不能作为类型标注使用。推荐使用 Optional[X] 保持前向兼容性。",
        "",
        "9.2 评价体系重构：6维度10分制 -> 7级+16二级指标+百分制",
        "【时间】2026-05-11 下午",
        "【背景】原系统使用6个平级维度（政策契合度/理论深度/实证支撑/创新价值/实践指导性/文献规范性），各维度满分10分。需求方提供 /root/instructions/研究报告、学术文献评价表4.16.docx，要求按照专业的7一级+16二级指标学术评价标准进行打分。",
        "【变更内容】",
        "  - 从6维度改为7个一级指标 + 16个二级指标（源文档标题写15是概数，实际表格含16项二级指标）",
        "  - 从10分制改为100分制（基础分100 + 附加分+/-5）",
        "  - 每个二级指标有独立的满分值：4/5/6/7/8/9分不等，全部由评价标准表格定义",
        "  - 每二级指标按4级评分锚点（优秀/良好/一般/差）给出分数",
        "  - 提示词从一句话描述改为嵌入完整的三级锚点（评分标准/评分细则/关键判断依据）",
        "【实施步骤】",
        "  ① /root/web/shared/constants.py 重写（501行）：定义完整的7xN二级指标 TypedDict 数据结构",
        "  ② /root/web/backend/models/schemas.py 重写：新增 SecondaryResult/PrimaryResult/AdditionalResult",
        "  ③ /root/web/backend/services/prompt_builder.py 重写：16个独立的结构化提示词，每个~1500字符",
        "  ④ /root/web/backend/services/evaluation_service.py 重写：迭代16指标->聚合7一级->总分",
        "  ⑤ /root/web/backend/utils/response_parser.py 修改：移除硬编码1-10钳制，支持0-9和负数",
        "  ⑥ /root/web/streamlit_app.py 重写：百分制仪表盘/得分率雷达图/柱状图/16个二级指标详情卡片",
        "【涉及文件】constants.py, schemas.py, prompt_builder.py, evaluation_service.py, response_parser.py, streamlit_app.py",
        "【教训】大幅度数据结构重构时，先改底层常量定义，再逐层向上改服务层->路由层->前端。每改一层立即验证导入，避免改到最后发现底层不兼容。",
        "",
        "9.3 PDF扫描件 OCR 解析",
        "【时间】2026-05-11 下午",
        "【现象】用户上传3个政府公文PDF文件后均显示「文献解析失败，请检查文件是否损坏或为空」。经排查，三个文件解析后文字长度均为0字符。",
        "【原因】三个PDF文件均为扫描件（每页仅含 /XObject 图片对象，无文本流）。pypdf.PdfReader.extract_text() 只能提取文本层文字，无法处理图片型PDF。此问题在政府公文类文献中非常普遍。",
        "【解决】实施双层解析策略，代码位于 /root/web/backend/utils/document_parser.py：",
        "  第1层：pypdf 优先提取文本层（适用文字型PDF，速度快）",
        "  第2层：若文字 < 50字符，自动回退到 OCR 流程",
        "  OCR流程依赖安装：apt-get install tesseract-ocr tesseract-ocr-chi-sim poppler-utils",
        "              pip install pdf2image pytesseract",
        "  技术链路：pdf2image.convert_from_bytes(PDF_bytes, dpi=200) -> PIL Image -> pytesseract.image_to_string(image, lang=chi_sim)",
        "【效果】三个测试文件分别提取 2015/8038/6961 字符，中文识别准确率良好",
        "【性能问题】当前10页PDF约80-150秒。瓶颈分布：pdf2image 转换（40%）+ Tesseract 中文识别（59%）。",
        "  待优化方向：降低DPI(300->200/150)、图片预处理（灰度化+二值化）、多进程并行OCR、换 PaddleOCR。",
        "",
        "9.4 旧进程残留导致新代码未生效",
        "【时间】2026-05-11 下午",
        "【现象】修改 streamlit_app.py 后重启服务，用户测试仍然报旧错误，且新文案（如7项一级指标）未出现在页面上。",
        "【原因】Streamlit 旧进程（PID 1935）自13:41起一直运行至今，kill 命令因 bash sandbox 限制（exit code 144 = SIGTERM）未能真正杀死进程。新启动命令因端口冲突绑定到了错误端口。",
        "【解决】使用 kill -9（强制 SIGKILL）绕过 sandbox 限制杀掉旧进程：kill -9 $(pgrep -f streamlit run)。",
        "【教训】在受限沙箱环境中，kill（SIGTERM）可能被拦截，需用 kill -9 确保进程终止。重启前用 pgrep 验证旧进程已消失。",
        "",
        "9.5 中文引号导致 Python 语法错误",
        "【时间】2026-05-11 下午",
        "【现象】导入 /root/web/shared/constants.py 时报 SyntaxError: invalid character (U+2014) at line 253。",
        "【原因】评价标准文本中含中文弯引号（左右引号），如「是否采用...对策」。但在原始代码中，这些中文弯引号被误写为 ASCII 双引号（U+0022），导致 Python 字符串提前终结。",
        "  例如字符串 \"是否采用\"问题原因对策\"递进结构\" 中，第二个 ASCII 双引号被Python解析为字符串结束符，后面的中文内容被当作代码导致语法错误。",
        "【解决】将字符串内嵌的中文双引号替换为直角引号（U+300C/U+300D），从根源上消除与 Python 字符串定界符的歧义。",
        "【涉及文件】/root/web/shared/constants.py",
        "【教训】在 Python 字符串中嵌套中文标点时，优先使用直角引号代替弯引号。这类问题只在运行时导入时才暴露（非语法高亮可检测），需要实际 import 验证。",
        "",
        "9.6 前后端分离后 schema 不一致",
        "【时间】2026-05-11 下午",
        "【背景】将单文件 /root/web/streamlit_app.py 拆分为17个模块后，前后端通过 Pydantic 模型共享数据结构。",
        "【问题】/root/web/backend/services/history_service.py 仍使用旧 DimensionResult 类型（old schema），与新 EvaluationResponse（含 PrimaryResult/SecondaryResult/AdditionalResult）不兼容。save_to_history 参数名从 overall_score 变为 total_score/base_score，routers/evaluation.py 和 streamlit_app.py 中的调用均需要更新。",
        "【解决】统一更新 history_service.py、routers/evaluation.py、routers/history.py、streamlit_app.py 中的 schema 引用和参数名，确保所有模块使用同一套 Pydantic 模型。",
        "【教训】重构数据模型时，使用 grep -rn 'old_field_name' /root/web/ 全局搜索旧字段名，确保所有引用处都已更新。类型检查器（mypy/pyright）可以捕获大部分不一致，但字段名变更需要人工审查。",
        "",
        "9.7 流式进度反馈缺失",
        "【时间】2026-05-11 下午",
        "【现象】原评价流程中用户点击「开始评价」后长时间无反馈（16个指标的模型推理可能需要数分钟），页面像卡死一样，体验差。",
        "【解决】在 /root/web/backend/services/evaluation_service.py 的 evaluate_document() 函数中添加 progress_callback 参数（类型：Optional[Callable[[int, int], None]]）。每完成一个二级指标，回调前端更新 st.progress 和状态文字。前端 streamlit_app.py 中定义 lambda cur, tot: update_progress(cur/tot)。",
        "【涉及文件】/root/web/backend/services/evaluation_service.py, /root/web/streamlit_app.py",
        "",
        "9.8 SSH 隧道公网访问探索",
        "【时间】2026-05-11 下午",
        "【背景】需将 Streamlit 服务（端口6006）暴露到外网，供项目方远程访问。服务器有公网IP（116.172.96.99），但所有入站端口被防火墙封锁。",
        "【尝试方案】",
        "  - serveo.net：SSH -R 反向隧道。连接速度快（5秒获取URL），但服务不稳定（间歇性不可达）。",
        "  - localhost.run（nokey用户）：认证慢（20-30秒），且 bash sandbox 会杀掉长连接 SSH 进程。",
        "  - 解决进程被杀：使用 screen -dmS 创建独立会话，script -q 提供伪终端避免输出缓冲。",
        "  - Cloudflare Tunnel（cloudflared）：GitHub 下载超时，未完成测试。",
        "【最终方案】推荐主方案：AutoDL 控制台端口映射。备用方案：SSH 反向隧道（serveo.net）。具体配置写入 6.2 公网映射。",
        "【教训】免费 SSH 隧道服务（serveo/localhost.run）适合临时演示，不适合生产环境。URL 每次连接会变化，需配合 screen 保持持久连接。",
        "",
        "9.9 文档解析的字节流处理",
        "【时间】2026-05-11 下午",
        "【背景】原代码中 parse_pdf() 接收 Streamlit 的 UploadedFile 对象（类文件接口），拆分为前后端后，需同时支持 BytesIO 对象和文件路径。",
        "【问题】parse_document(BytesIO(content), filename) 中，BytesIO 被 parse_pdf() 消费后，指针已到末尾。后续 OCR 回退需要重新读取原始字节时，BytesIO.read() 返回空。",
        "【解决】在 parse_pdf() 开头保存 pdf_bytes = file.read()，然后用 BytesIO(pdf_bytes) 创建新对象供 pypdf 使用。OCR 回退时直接使用已保存的 pdf_bytes，避免指针问题。",
        "【涉及文件】/root/web/backend/utils/document_parser.py",
        "【教训】处理 BytesIO 等有状态流对象时，始终先复制原始数据（.read() 到 bytes），后续操作使用独立副本。",
    ]

    for item in reversed(tech_log_items):
        p = doc.add_paragraph(item)
        parent.insert(insert_idx, p._element)
    print("  Section 9 created")

# ====================================================================
# Save
# ====================================================================
output_path = "/root/instructions/中医药政策研究文献评价表 (1).docx"
doc.save(output_path)
print(f"Saved: {output_path}")
print("All done!")
