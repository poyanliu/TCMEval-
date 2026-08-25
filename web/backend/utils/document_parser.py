"""Document parsing utilities — PDF (text + OCR) and DOCX extraction."""

import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import BytesIO
from typing import Optional

from pypdf import PdfReader
from docx import Document as DocxDocument

from backend.config import OCR_DPI, OCR_LANG, OCR_WORKERS, OCR_PREPROCESS, OCR_PSM, OCR_MAX_IMAGE_PX

logger = logging.getLogger(__name__)

# Lazy imports for OCR (heavy, only needed for scanned PDFs)
_pdf2image = None
_pytesseract = None
_Image = None  # PIL Image, for preprocessing


def _get_pdf2image():
    global _pdf2image
    if _pdf2image is None:
        import pdf2image as p2i
        _pdf2image = p2i
    return _pdf2image


def _get_pytesseract():
    global _pytesseract
    if _pytesseract is None:
        import pytesseract as pts
        _pytesseract = pts
    return _pytesseract


def _get_Image():
    global _Image
    if _Image is None:
        from PIL import Image as _Img
        _Image = _Img
    return _Image


# ── Image preprocessing ────────────────────────────────────────────
def _preprocess_image(image) -> object:
    """Grayscale + binarize with Otsu threshold for faster/better OCR.

    Returns the original image unchanged if OCR_PREPROCESS is disabled.
    """
    if not OCR_PREPROCESS:
        return image

    Image = _get_Image()
    img = image.convert("L")  # grayscale
    # Otsu binarization: split into pure black/white
    threshold = _otsu_threshold(img)
    img = img.point(lambda p: 0 if p < threshold else 255)
    return img


def _otsu_threshold(img) -> int:
    """Compute Otsu's threshold for a grayscale PIL Image."""
    hist = img.histogram()
    total = sum(hist)
    if total == 0:
        return 128

    sum_b = 0
    w_b = 0
    maximum = 0.0
    sum_total = sum(i * hist[i] for i in range(256))
    threshold = 128

    for i in range(256):
        w_b += hist[i]
        if w_b == 0:
            continue
        w_f = total - w_b
        if w_f == 0:
            break
        sum_b += i * hist[i]
        mean_b = sum_b / w_b
        mean_f = (sum_total - sum_b) / w_f
        between = w_b * w_f * (mean_b - mean_f) ** 2
        if between > maximum:
            maximum = between
            threshold = i

    return threshold


# ── PDF parsing ────────────────────────────────────────────────────
def parse_pdf(file, enable_ocr: bool = True) -> str:
    """Extract text from an uploaded PDF file.

    First attempts pypdf (fast, works for text-based PDFs). If the
    extracted text is empty or near-empty, falls back to OCR via
    pytesseract (for scanned/image-based PDFs).
    """
    if hasattr(file, 'read'):
        pdf_bytes = file.read()
        file = BytesIO(pdf_bytes)
    else:
        pdf_bytes = None

    # ── Step 1: Try pypdf (fast, works for text PDFs) ────────────
    reader = PdfReader(file)
    texts: list[str] = []
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            texts.append(extracted.strip())

    text = "\n".join(texts)

    # ── Step 2: Fall back to OCR if pypdf got nothing ────────────
    if (not text or len(text.strip()) < 50) and enable_ocr and pdf_bytes:
        logger.info("PDF appears to be scanned (no extractable text), running OCR...")
        try:
            ocr_text = _ocr_pdf(pdf_bytes)
            if ocr_text:
                return ocr_text
        except Exception as exc:
            logger.warning("OCR fallback failed: %s", exc)

    return text


def _ocr_single_page(args: tuple) -> tuple[int, str]:
    """OCR a single page image. Top-level function for picklability."""
    i, image = args
    pytesseract = _get_pytesseract()
    processed = _preprocess_image(image)
    # Resize if image exceeds max dimension (preserving aspect ratio)
    w, h = processed.size
    max_px = max(w, h)
    if max_px > OCR_MAX_IMAGE_PX:
        scale = OCR_MAX_IMAGE_PX / max_px
        processed = processed.resize(
            (int(w * scale), int(h * scale)),
            _get_Image().LANCZOS,
        )
    # PSM 6: assume uniform block of text — faster than default PSM 3
    page_text = pytesseract.image_to_string(
        processed, lang=OCR_LANG, config=f"--psm {OCR_PSM}"
    )
    return i, page_text.strip()


def _ocr_pdf(pdf_bytes: bytes) -> str:
    """OCR all pages of a scanned PDF using Tesseract with multithreading.

    Converts pages lazily one at a time and OCRs them in parallel using
    a thread pool (Tesseract releases the GIL for its C++ core).
    Image preprocessing (grayscale + Otsu binarization) is applied to
    each page before OCR for better speed and accuracy.
    """
    pdf2image = _get_pdf2image()

    # Convert to images first (pdf2image is single-threaded C-backed work)
    logger.info("Converting PDF to images at %d DPI...", OCR_DPI)
    images = pdf2image.convert_from_bytes(pdf_bytes, dpi=OCR_DPI)
    total = len(images)
    logger.info("PDF has %d pages, OCRing with %d workers", total, OCR_WORKERS)

    # Sort by page index to preserve order
    results: dict[int, str] = {}

    with ThreadPoolExecutor(max_workers=OCR_WORKERS) as executor:
        futures = {
            executor.submit(_ocr_single_page, (i, img)): i
            for i, img in enumerate(images)
        }
        for future in as_completed(futures):
            i, page_text = future.result()
            if page_text:
                results[i] = page_text
            if len(results) % 5 == 0 or len(results) == total:
                logger.info("OCR progress: %d/%d pages", len(results), total)

    return "\n".join(results[i] for i in sorted(results) if results[i])


# ── DOCX parsing ───────────────────────────────────────────────────
def parse_docx(file) -> str:
    """Extract full text from an uploaded DOCX file."""
    doc = DocxDocument(BytesIO(file.read()))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)


# ── Unified entry point ────────────────────────────────────────────
SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def parse_document(file, filename: str, enable_vision: bool = True) -> tuple[str, str]:
    """Parse an uploaded document, auto-detecting format from filename.

    Returns enriched text that includes:
      - Body text extracted from the document
      - Tables converted to markdown
      - Image descriptions from the vision model (if enable_vision=True)
    """
    ext = filename.lower()

    # Extract body text
    if ext.endswith(".pdf"):
        body_text = parse_pdf(file)
    elif ext.endswith(".docx"):
        body_text = parse_docx(file)
    else:
        raise ValueError(
            f"不支持的文件格式 '{filename}'，请上传 PDF 或 DOCX 文件"
        )

    # Reset file position for subsequent extraction passes
    if hasattr(file, 'seek'):
        file.seek(0)

    # Extract tables
    table_mds = extract_tables(file, filename)
    tables_section = ""
    if table_mds:
        tables_section = (
            "\n\n【文献中的表格数据】\n"
            + "\n---\n".join(table_mds)
        )

    # Extract images and describe with vision model
    images_section = ""
    if enable_vision:
        try:
            from backend.services.vision_client import describe_images
            images = extract_images(file, filename)
            if images:
                descriptions = describe_images(images, context=filename)
                if descriptions:
                    parts = []
                    for d in descriptions:
                        parts.append(
                            f"图片{d['index'] + 1}（{d['type']}）：{d['description']}"
                        )
                    images_section = (
                        "\n\n【文献中的图表内容（视觉模型识别）】\n"
                        + "\n".join(parts)
                    )
        except Exception:
            pass  # Vision is best-effort; body text + tables are sufficient

    # Combine
    enriched = body_text
    if tables_section:
        enriched += tables_section
    if images_section:
        enriched += images_section

    return filename, enriched


def detect_format(filename: str) -> Optional[str]:
    """Return MIME type for a supported file, or None."""
    for ext, mime in SUPPORTED_EXTENSIONS.items():
        if filename.lower().endswith(ext):
            return mime
    return None


# ── Table extraction ──────────────────────────────────────────────
def _table_to_markdown(rows: list[list[str]]) -> str:
    """Convert a 2D list of cell strings to a markdown table."""
    if not rows:
        return ""
    max_cols = max(len(row) for row in rows)
    # Normalize
    padded = [row + [""] * (max_cols - len(row)) for row in rows]
    lines = []
    header = "| " + " | ".join(str(c) if c else "" for c in padded[0]) + " |"
    lines.append(header)
    sep = "| " + " | ".join("---" for _ in range(max_cols)) + " |"
    lines.append(sep)
    for row in padded[1:]:
        lines.append("| " + " | ".join(str(c) if c else "" for c in row) + " |")
    return "\n".join(lines)


def extract_tables_pdf(file) -> list[str]:
    """Extract tables from a PDF using PyMuPDF (fitz).

    Returns list of markdown table strings.
    """
    import fitz
    if hasattr(file, 'read'):
        pdf_bytes = file.read()
        file.seek(0)
    else:
        pdf_bytes = file

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    tables: list[str] = []
    try:
        for page in doc:
            tabs = page.find_tables()
            if tabs and tabs.tables:
                for table in tabs.tables:
                    md = _table_to_markdown(table.extract())
                    if md.strip():
                        tables.append(md)
    finally:
        doc.close()
    return tables


def extract_tables_docx(file) -> list[str]:
    """Extract tables from a DOCX using python-docx.

    Returns list of markdown table strings.
    """
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    doc = DocxDocument(file)
    tables: list[str] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(cells)
        if rows:
            md = _table_to_markdown(rows)
            if md.strip():
                tables.append(md)
    return tables


def extract_tables(file, filename: str) -> list[str]:
    """Extract tables from PDF or DOCX, returning markdown table strings."""
    ext = filename.lower()
    if ext.endswith(".pdf"):
        return extract_tables_pdf(file)
    elif ext.endswith(".docx"):
        return extract_tables_docx(file)
    return []


# ── Image extraction ───────────────────────────────────────────────
def _resize_image(img, max_dim: int = 1024) -> object:
    """Resize a PIL Image so its largest dimension <= max_dim."""
    w, h = img.size
    max_px = max(w, h)
    if max_px <= max_dim:
        return img
    scale = max_dim / max_px
    return img.resize((int(w * scale), int(h * scale)), _get_Image().LANCZOS)


def extract_images_pdf(file, max_images: int = 10) -> list:
    """Extract embedded images from a PDF using PyMuPDF.

    Returns list of PIL Images (max max_images, skipping duplicates by size).
    """
    import fitz

    if hasattr(file, 'read'):
        pdf_bytes = file.read()
        file.seek(0)
    else:
        pdf_bytes = file

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images: list = []
    seen_sizes: set[tuple[int, int]] = set()
    Image = _get_Image()

    try:
        for page in doc:
            if len(images) >= max_images:
                break
            img_list = page.get_images(full=True)
            for img_info in img_list:
                if len(images) >= max_images:
                    break
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    img_bytes = base_image["image"]
                    img = Image.open(BytesIO(img_bytes))
                    img = img.convert("RGB")
                    # Skip very small images (icons, logos < 100px in both dims)
                    if img.width < 100 and img.height < 100:
                        continue
                    # Skip duplicates by size
                    size_key = (img.width, img.height)
                    if size_key in seen_sizes:
                        continue
                    seen_sizes.add(size_key)
                    img = _resize_image(img)
                    images.append(img)
                except Exception:
                    continue
    finally:
        doc.close()
    return images


def extract_images_docx(file, max_images: int = 10) -> list:
    """Extract embedded images from a DOCX file.

    Returns list of PIL Images.
    """
    from docx import Document as DocxDocument
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc = DocxDocument(file)
    images: list = []
    Image = _get_Image()

    for rel in doc.part.rels.values():
        if len(images) >= max_images:
            break
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                img = Image.open(BytesIO(img_bytes))
                img = img.convert("RGB")
                if img.width < 100 and img.height < 100:
                    continue
                img = _resize_image(img)
                images.append(img)
            except Exception:
                continue
    return images


def extract_images(file, filename: str, max_images: int = 10) -> list:
    """Extract embedded images from PDF or DOCX. Returns list of PIL Images."""
    ext = filename.lower()
    if ext.endswith(".pdf"):
        return extract_images_pdf(file, max_images)
    elif ext.endswith(".docx"):
        return extract_images_docx(file, max_images)
    return []


def truncate_text(text: str, max_chars: int = 6000) -> str:
    """Truncate document text for LLM context window."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "…[内容截断]"
