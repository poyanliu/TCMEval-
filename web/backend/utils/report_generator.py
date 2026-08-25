"""Report generator — DOCX and PDF exports for evaluation results.

Generates professionally formatted evaluation reports with:
  - DOCX: styled headings, score tables, evidence blocks (python-docx)
  - PDF:  Chinese typography, structured layout (fpdf2)
  - TXT:  plain-text fallback (included for convenience)

All generators consume the same EvaluationResponse + primary indicator config.
"""

from __future__ import annotations

import io
import os
from datetime import datetime
from typing import Optional

from backend.models.schemas import EvaluationResponse, PrimaryResult, SecondaryResult
from shared.constants import PRIMARY_INDICATORS, OVERALL_THRESHOLDS

# ═══════════════════════════════════════════════════════════════════════
# Shared helpers
# ═══════════════════════════════════════════════════════════════════════

_TITLE = "中医药政策文献智能评价报告"
_SUBTITLE = "基于 GLM-4-9B 大模型 | 7项一级指标 · 16项二级指标（部分可选）| 百分制"


def _score_ratio_label(score: int, max_score: int) -> str:
    """Return a quality label based on score/max ratio."""
    if max_score == 0:
        return "N/A"
    ratio = score / max_score
    if ratio >= 0.85:
        return "优秀"
    elif ratio >= 0.70:
        return "良好"
    elif ratio >= 0.50:
        return "一般"
    else:
        return "待提升"


def _build_summary_text(response: EvaluationResponse) -> str:
    """Build plain-text summary for DOCX/PDF content."""
    excluded = getattr(response, 'excluded_indicators', [])
    max_total = 110 if response.additional_results else 100
    lines = [
        f"文献名称：{response.doc_name}",
        f"评价时间：{response.timestamp}",
        f"总分：{response.total_score} / {max_total:.0f}（基础分 {response.base_score} / 100",
    ]
    bonus_total = sum(a.score for a in response.additional_results) if response.additional_results else 0
    if bonus_total != 0:
        lines[-1] += f"，附加分 {bonus_total:+d}"
    lines[-1] += "）"
    if excluded:
        lines.append(f"已跳过指标：{', '.join(excluded)}（文献不含对应内容，已等比缩放至百分制）")
    lines.append("")

    lines.append("一级指标得分：")
    for p in response.primary_results:
        pct = (p.score / p.weight * 100) if p.weight > 0 else 0
        lines.append(f"  {p.id}、{p.name}: {p.score}/{p.weight} ({pct:.0f}%)")

    lines.append("")
    lines.append("二级指标明细：")
    for p in response.primary_results:
        for r in p.secondary_results:
            lines.append(f"  {r.id} {r.name}: {r.score}/{r.max_score} — {r.comment}")

    lines.append("")
    lines.append("综合评价：")
    lines.append(response.overall_comment)

    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════════
# DOCX generator
# ═══════════════════════════════════════════════════════════════════════

def generate_docx(response: EvaluationResponse) -> bytes:
    """Generate a styled .docx evaluation report.

    Returns:
        DOCX file content as bytes (ready for download).
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # ── Title ─────────────────────────────────────────────────────
    title = doc.add_heading(_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1a, 0x4d, 0x2c)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run(_SUBTITLE)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Meta info ─────────────────────────────────────────────────
    doc.add_paragraph()
    excluded = getattr(response, 'excluded_indicators', [])
    max_total = 110 if response.additional_results else 100
    meta_cells = [
        ("文献名称", response.doc_name),
        ("评价时间", response.timestamp),
        ("总分", f"{response.total_score} / {max_total:.0f}"),
        ("基础分", f"{response.base_score} / 100"),
    ]
    if excluded:
        meta_cells.append(("已跳过指标", ", ".join(excluded) + "（已等比缩放至百分制）"))
    meta_table = doc.add_table(rows=len(meta_cells), cols=2)
    meta_table.style = "Table Grid"
    for i, (label, value) in enumerate(meta_cells):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        for cell in meta_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        # Bold the label column
        for run in meta_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    for a in (response.additional_results or []):
        if a.score != 0:
            row = meta_table.add_row()
            row.cells[0].text = f"附加·{a.name}"
            row.cells[1].text = f"{a.score:+d}"

    doc.add_paragraph()

    # ── Primary indicators overview table ─────────────────────────
    doc.add_heading("一、一级指标得分总览", level=1)
    overview_table = doc.add_table(rows=1, cols=5)
    overview_table.style = "Table Grid"
    overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["编号", "一级指标", "权重", "得分", "得分率"]
    hdr_cells = overview_table.rows[0].cells
    for i, hdr in enumerate(headers):
        hdr_cells[i].text = hdr
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9.5)

    for p in response.primary_results:
        row = overview_table.add_row()
        pct = (p.score / p.weight * 100) if p.weight > 0 else 0
        values = [p.id, p.name, str(p.weight), str(p.score), f"{pct:.0f}%"]
        for i, val in enumerate(values):
            row.cells[i].text = val
            for run in row.cells[i].paragraphs[0].runs:
                run.font.size = Pt(9.5)

    doc.add_paragraph()

    # ── Detailed indicator results ────────────────────────────────
    doc.add_heading("二、分指标评价详情", level=1)
    for p in response.primary_results:
        doc.add_heading(f"{p.id}、{p.name}（{p.score}/{p.weight} 分）", level=2)

        for r in p.secondary_results:
            ratio = (r.score / r.max_score * 100) if r.max_score > 0 else 0
            label = _score_ratio_label(r.score, r.max_score)

            # Indicator header
            para = doc.add_paragraph()
            run_ind = para.add_run(f"■ {r.id} {r.name}")
            run_ind.bold = True
            run_ind.font.size = Pt(11)

            run_score = para.add_run(f"    {r.score} / {r.max_score} 分 ({ratio:.0f}%)  [{label}]")
            run_score.font.size = Pt(10)
            run_score.font.color.rgb = RGBColor(
                0x1b, 0x5e, 0x20,
            ) if ratio >= 70 else RGBColor(0xe5, 0x39, 0x35)

            # Evidence
            ev_para = doc.add_paragraph()
            ev_run = ev_para.add_run(f"证据：{r.evidence}")
            ev_run.font.size = Pt(9.5)
            ev_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Comment
            cm_para = doc.add_paragraph()
            cm_run = cm_para.add_run(f"评语：{r.comment}")
            cm_run.font.size = Pt(9.5)
            cm_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    # ── Overall assessment ────────────────────────────────────────
    doc.add_heading("三、综合评价", level=1)
    overall_para = doc.add_paragraph(response.overall_comment)

    active_bonuses = [a for a in (response.additional_results or []) if a.score != 0]
    if active_bonuses:
        doc.add_heading("四、附加项评价", level=1)
        for a in active_bonuses:
            add_para = doc.add_paragraph()
            add_para.add_run(f"{a.name}：").bold = True
            add_para.add_run(f"得分 {a.score:+d} 分 — {a.comment}")

    # ── Footer ────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("— 本报告由中医药政策文献智能评价系统自动生成 —")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Serialize ─────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════
# PDF generator
# ═══════════════════════════════════════════════════════════════════════

# fpdf2 cannot handle .ttc (TrueType Collection) files correctly — glyphs
# render as garbled text in many viewers.  We extract the Simplified Chinese
# weight from each TTC as a standalone .ttf on first use.
_CJK_FONT_REGULAR: str | None = None
_CJK_FONT_BOLD: str | None = None
_TTC_EXTRACT_DIR = "/tmp/cjk_fonts"


def _extract_font_from_ttc(ttc_path: str, output_name: str, index: int = 0) -> str:
    """Extract a single font from a TTC using fontTools, return .ttf path."""
    from fontTools.ttLib import TTCollection
    os.makedirs(_TTC_EXTRACT_DIR, exist_ok=True)
    out_path = os.path.join(_TTC_EXTRACT_DIR, output_name)
    if os.path.exists(out_path):
        return out_path
    ttc = TTCollection(ttc_path)
    font = ttc[index]
    font.save(out_path)
    return out_path


_FONT_SPECS = [
    # Prefer TrueType (glyf) fonts — fpdf2 embeds them as TrueType, no mismatch.
    ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", 0, "WenQuanYiMicroHei.ttf"),
    # Fall back to CFF-based Noto CJK (works but some viewers flag a warning).
    ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", 2, "NotoSansCJKsc-Regular.ttf"),
    ("/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc", 2, "NotoSerifCJKsc-Regular.ttf"),
    ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", 2, "NotoSansCJKsc-Regular.ttf"),
]


def _get_cjk_font_regular() -> str:
    global _CJK_FONT_REGULAR
    if _CJK_FONT_REGULAR is None:
        for path, idx, name in _FONT_SPECS:
            if os.path.exists(path):
                _CJK_FONT_REGULAR = _extract_font_from_ttc(path, name, idx)
                break
        else:
            _CJK_FONT_REGULAR = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    return _CJK_FONT_REGULAR


def _get_cjk_font_bold() -> str:
    global _CJK_FONT_BOLD
    if _CJK_FONT_BOLD is None:
        # Reuse the same TrueType font; fpdf2 simulates bold via text-stroking.
        # Noto Sans CJK Bold is CFF-based and would cause mixed-font warnings.
        _CJK_FONT_BOLD = _get_cjk_font_regular()
    return _CJK_FONT_BOLD


def generate_pdf(response: EvaluationResponse) -> bytes:
    """Generate a styled PDF evaluation report with Chinese typography.

    Uses fpdf2 with a CJK TrueType font.  The layout mirrors the DOCX
    report: title page → overview table → detail blocks → overall comment.

    Returns:
        PDF file content as bytes (ready for download).
    """
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")

    # Register extracted standalone CJK fonts (not TTC)
    pdf.add_font("CJK", "", _get_cjk_font_regular())
    pdf.add_font("CJK", "B", _get_cjk_font_bold())

    pdf.set_auto_page_break(auto=True, margin=18)

    # ── Page 1: Title & overview ──────────────────────────────────
    pdf.add_page()

    # Title
    pdf.set_font("CJK", "B", 22)
    pdf.set_text_color(0x1a, 0x4d, 0x2c)
    pdf.cell(0, 14, _TITLE, align="C")
    pdf.ln(12)

    # Subtitle
    pdf.set_font("CJK", "", 9)
    pdf.set_text_color(0x66, 0x66, 0x66)
    pdf.cell(0, 6, _SUBTITLE, align="C")
    pdf.ln(12)

    # Separator line
    y_before = pdf.get_y()
    pdf.set_draw_color(0x2c, 0x77, 0x44)
    pdf.set_line_width(0.4)
    pdf.line(25, y_before, 185, y_before)
    pdf.ln(8)

    # Meta info
    pdf.set_font("CJK", "", 10)
    pdf.set_text_color(0x33, 0x33, 0x33)
    excluded = getattr(response, 'excluded_indicators', [])
    max_total = 110 if response.additional_results else 100
    meta_lines = [
        f"文献名称：{response.doc_name}",
        f"评价时间：{response.timestamp}",
        f"总  分：{response.total_score} / {max_total:.0f}（基础分 {response.base_score} / 100",
    ]
    bonus_total = sum(a.score for a in response.additional_results) if response.additional_results else 0
    if bonus_total != 0:
        meta_lines[-1] += f"，附加分 {bonus_total:+d}"
    meta_lines[-1] += "）"
    if excluded:
        meta_lines.append(f"已跳过指标：{', '.join(excluded)}（已等比缩放至百分制）")
    for a in (response.additional_results or []):
        if a.score != 0:
            meta_lines.append(f"附加项·{a.name}：{a.score:+d} 分")

    for line in meta_lines:
        pdf.cell(0, 7, line)
        pdf.ln(7)
    pdf.ln(4)

    # ── Section 1: Primary indicators overview ────────────────────
    pdf.set_font("CJK", "B", 14)
    pdf.set_text_color(0x1a, 0x4d, 0x2c)
    pdf.cell(0, 9, "一、一级指标得分总览")
    pdf.ln(12)

    # Table header
    col_widths = [12, 56, 16, 16, 20]
    pdf.set_font("CJK", "B", 9)
    pdf.set_fill_color(0xe8, 0xf5, 0xe9)
    headers = ["编号", "一级指标", "权重", "得分", "得分率"]
    for i, (hdr, w) in enumerate(zip(headers, col_widths)):
        pdf.cell(w, 8, hdr, border=1, fill=True, align="C")
    pdf.ln()

    # Table rows
    pdf.set_font("CJK", "", 9)
    for p in response.primary_results:
        pct = (p.score / p.weight * 100) if p.weight > 0 else 0
        values = [p.id, p.name, str(p.weight), str(p.score), f"{pct:.0f}%"]
        for i, (val, w) in enumerate(zip(values, col_widths)):
            align = "C" if i in (0, 2, 3, 4) else "L"
            # Use short name for table
            display_val = val if i != 1 else val[:12]
            pdf.cell(w, 7, display_val, border=1, align=align)
        pdf.ln()
    pdf.ln(6)

    # ── Section 2: Detailed indicator results ─────────────────────
    pdf.set_font("CJK", "B", 14)
    pdf.set_text_color(0x1a, 0x4d, 0x2c)
    pdf.cell(0, 9, "二、分指标评价详情")
    pdf.ln(12)

    for p in response.primary_results:
        # Primary header
        pdf.set_font("CJK", "B", 11)
        pdf.set_text_color(0x1a, 0x4d, 0x2c)
        pct = (p.score / p.weight * 100) if p.weight > 0 else 0
        pdf.cell(0, 7, f"{p.id}、{p.name}（{p.score}/{p.weight} 分, {pct:.0f}%）")
        pdf.ln(9)

        for r in p.secondary_results:
            ratio = (r.score / r.max_score * 100) if r.max_score > 0 else 0
            label = _score_ratio_label(r.score, r.max_score)

            # Check if we need a page break (at least 30mm remaining)
            if pdf.get_y() > 255:
                pdf.add_page()

            # Indicator header
            pdf.set_font("CJK", "B", 10)
            pdf.set_text_color(0x22, 0x22, 0x22)
            pdf.cell(0, 6, f"■ {r.id} {r.name}    {r.score}/{r.max_score} 分 ({ratio:.0f}%)  [{label}]")
            pdf.ln(7)

            # Evidence
            pdf.set_font("CJK", "", 9)
            pdf.set_text_color(0x55, 0x55, 0x55)
            evidence_text = f"证据：{r.evidence}"
            pdf.multi_cell(0, 5.5, evidence_text)
            pdf.ln(1)

            # Comment
            pdf.set_font("CJK", "", 9)
            pdf.set_text_color(0x44, 0x44, 0x44)
            pdf.cell(0, 5.5, f"评语：{r.comment}")
            pdf.ln(7)

    # ── Section 3: Overall assessment ─────────────────────────────
    if pdf.get_y() > 230:
        pdf.add_page()
    pdf.ln(4)
    pdf.set_draw_color(0x2c, 0x77, 0x44)
    pdf.set_line_width(0.4)
    y = pdf.get_y()
    pdf.line(25, y, 185, y)
    pdf.ln(6)

    pdf.set_font("CJK", "B", 14)
    pdf.set_text_color(0x1a, 0x4d, 0x2c)
    pdf.cell(0, 9, "三、综合评价")
    pdf.ln(12)

    pdf.set_font("CJK", "", 10)
    pdf.set_text_color(0x33, 0x33, 0x33)
    pdf.multi_cell(0, 6.5, response.overall_comment)

    # ── Additional item ────────────────────────────────────────────
    active_bonuses = [a for a in (response.additional_results or []) if a.score != 0]
    if active_bonuses:
        pdf.ln(6)
        pdf.set_font("CJK", "B", 14)
        pdf.set_text_color(0x1a, 0x4d, 0x2c)
        pdf.cell(0, 9, "四、附加项评价")
        pdf.ln(12)
        pdf.set_font("CJK", "", 10)
        pdf.set_text_color(0x33, 0x33, 0x33)
        for a in active_bonuses:
            pdf.cell(0, 6.5, f"{a.name}：得分 {a.score:+d} 分 — {a.comment}")
            pdf.ln(8)

    # ── Footer ────────────────────────────────────────────────────
    pdf.ln(10)
    pdf.set_font("CJK", "", 7)
    pdf.set_text_color(0x99, 0x99, 0x99)
    pdf.cell(0, 5, "— 本报告由中医药政策文献智能评价系统自动生成 —", align="C")

    # ── Serialize ─────────────────────────────────────────────────
    return bytes(pdf.output())


# ═══════════════════════════════════════════════════════════════════════
# Convenience: generate all formats
# ═══════════════════════════════════════════════════════════════════════

def generate_report(
    response: EvaluationResponse,
    fmt: str = "docx",
) -> bytes:
    """Generate an evaluation report in the specified format.

    Args:
        response: Completed EvaluationResponse.
        fmt: One of "docx", "pdf", "txt".

    Returns:
        Report content as bytes.
    """
    if fmt == "docx":
        return generate_docx(response)
    elif fmt == "pdf":
        return generate_pdf(response)
    elif fmt == "txt":
        return _build_summary_text(response).encode("utf-8")
    else:
        raise ValueError(f"Unsupported format: {fmt}")


def get_mime_type(fmt: str) -> str:
    """Return the MIME type for a report format."""
    return {
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf": "application/pdf",
        "txt": "text/plain",
    }.get(fmt, "application/octet-stream")


def get_file_extension(fmt: str) -> str:
    """Return the file extension for a report format."""
    return {"docx": ".docx", "pdf": ".pdf", "txt": ".txt"}.get(fmt, "")
