import os
import json
import sqlite3
from datetime import datetime
from pathlib import Path

import docx
import fitz  # PyMuPDF
from flask import Flask, request, jsonify, send_from_directory, g
from openai import OpenAI
from werkzeug.utils import secure_filename

# Load .env file (if present) — keeps secrets out of source control
_ENV_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
if os.path.exists(_ENV_FILE):
    with open(_ENV_FILE, "r", encoding="utf-8") as _f:
        for _line in _f:
            _line = _line.strip()
            if _line and not _line.startswith("#") and "=" in _line:
                _k, _, _v = _line.partition("=")
                _k, _v = _k.strip(), _v.strip().strip("\"'")
                if _k and _k not in os.environ:
                    os.environ[_k] = _v

app = Flask(__name__, static_folder="static", template_folder="templates")
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024  # 32MB
app.config["UPLOAD_FOLDER"] = "/root/EcoEval/uploads"

ALLOWED_EXTENSIONS = {"docx", "pdf", "doc"}

# DeepSeek API configuration
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
DEEPSEEK_MODEL = os.environ.get("DEEPSEEK_MODEL", "deepseek-chat")

DB_PATH = "/root/EcoEval/evaluations.db"


def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
    return g.db


@app.teardown_appcontext
def close_db(exception):
    db = g.pop("db", None)
    if db is not None:
        db.close()


def init_db():
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS evaluations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                original_filename TEXT NOT NULL,
                file_type TEXT NOT NULL,
                file_content TEXT NOT NULL,
                file_data BLOB,
                score_json TEXT NOT NULL,
                total_score REAL NOT NULL,
                dimension_scores TEXT NOT NULL,
                created_at TEXT NOT NULL
            )
        """)


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(filepath, ext):
    if ext == "docx":
        doc = docx.Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                tables_text.append(row_text)
        return "\n\n".join(paragraphs + tables_text)

    elif ext == "pdf":
        pdf = fitz.open(filepath)
        text = ""
        for page in pdf:
            text += page.get_text()
        pdf.close()
        return text

    return ""


EVALUATION_CRITERIA = """
你是一位中医治未病卫生经济学评价专家。请根据以下评价指标体系，对用户上传的研究文档进行评分。

## 评分规则
- 每个三级指标满分100分，根据文档中是否包含该指标相关数据及详细程度打分：
  - 0分：完全不涉及
  - 1-30分：仅提及概念，无具体数据
  - 31-60分：有部分数据但不完整
  - 61-85分：数据较完整，稍有缺项
  - 86-100分：数据完整，描述清晰
- 每个指标得分 × 该指标权重系数（权重% / 100） = 加权得分
- 总分 = 所有加权得分之和（满分100分）

## 评价指标体系

### 一、人群基线与中医体质基线指标（总权重15%）
| 序号 | 二级指标 | 三级指标 | 指标描述 | 权重 |
|------|---------|---------|---------|------|
| 1 | 通用人口学基线 | 年龄 | 入组时实际年龄（岁） | 0.6% |
| 2 | 通用人口学基线 | 性别 | 男/女 | 0.6% |
| 3 | 通用人口学基线 | 职业 | 公职人员/企业职工/自由职业/退休/学生/其他 | 0.6% |
| 4 | 通用人口学基线 | 婚姻状况 | 未婚/已婚/离异/丧偶 | 0.6% |
| 5 | 通用人口学基线 | 文化水平 | 小学及以下/中学/中专/大专/本科及以上 | 0.6% |
| 6 | 通用人口学基线 | 家庭年收入 | 收入分层数据 | 0.6% |
| 7 | 通用人口学基线 | 民族 | 汉族/少数民族 | 0.4% |
| 8 | 通用人口学基线 | 医保类型 | 职工医保/居民医保/新农合/商业保险/无医保 | 0.6% |
| 9 | 通用人口学基线 | 既往慢病史 | 高血压/糖尿病/冠心病/脑梗/高血脂/其他 | 0.8% |
| 10 | 通用人口学基线 | 身高 | cm | 0.4% |
| 11 | 通用人口学基线 | 体重 | kg | 0.4% |
| 12 | 通用人口学基线 | BMI指数 | 体重(kg)/身高²(m²) | 0.6% |
| 13 | 通用人口学基线 | 血压 | 收缩压/舒张压（mmHg） | 0.8% |
| 14 | 通用人口学基线 | 空腹血糖 | mmol/L | 0.8% |
| 15 | 通用人口学基线 | 糖化血红蛋白 | % | 0.8% |
| 16 | 通用人口学基线 | 血脂四项 | 总胆固醇、低密度脂蛋白、高密度脂蛋白、甘油三酯 | 0.8% |
| 17 | 通用人口学基线 | 入组前1年直接医疗总支出 | 药品、检查、住院、诊疗等费用总和（元） | 1.0% |
| 18 | 中医特色体质基线 | 中医九种体质判定 | 平和质/气虚质/阳虚质/阴虚质/痰湿质/湿热质/血瘀质/气郁质/特禀质 | 1.2% |
| 19 | 中医特色体质基线 | 中医证候总积分 | 疲乏、失眠、畏寒、胸闷、腹胀、肢体困重等单项证候评分总和 | 1.2% |
| 20 | 中医特色体质基线 | 单项中医证候积分 | 各单项证候具体评分 | 1.0% |
| 21 | 中医特色体质基线 | 中医健康素养水平 | 中医养生知识掌握程度问卷得分 | 0.8% |
| 22 | 中医特色体质基线 | 治未病认知程度 | 完全不了解/了解较少/一般了解/比较了解/非常了解 | 0.8% |

### 二、治未病干预直接成本指标（总权重20%）
| 序号 | 二级指标 | 三级指标 | 指标描述 | 权重 |
|------|---------|---------|---------|------|
| 23 | 中医综合调理干预成本 | 中药膏方/汤药费 | 总费用、医保报销、个人自付（元） | 2.2% |
| 24 | 中医综合调理干预成本 | 中医外治服务费 | 艾灸、拔罐、刮痧、穴位贴敷、推拿、针灸、耳穴压豆总费用 | 2.2% |
| 25 | 中医综合调理干预成本 | 中医养生功法指导费 | 八段锦、太极拳、五禽戏、导引等教学服务费 | 1.8% |
| 26 | 中医综合调理干预成本 | 中医体质辨识费 | 体质辨识评估、建档费用 | 1.8% |
| 27 | 中医综合调理干预成本 | 健康随访管理费 | 季度随访、体质复评、健康咨询服务费 | 2.0% |
| 28 | 中医综合调理干预成本 | 食疗养生指导费用 | 药食同源食材、药膳调理配套支出 | 2.0% |
| 29 | 西医配套及其他成本 | 常规药品费用 | 降压、降糖、调脂等慢病西药 | 1.8% |
| 30 | 西医配套及其他成本 | 检查检验费用 | 血压、血糖、血脂、影像学等年度筛查费用 | 1.6% |
| 31 | 西医配套及其他成本 | 住院诊疗费用 | 慢病急性发作住院、并发症住院 | 1.8% |
| 32 | 西医配套及其他成本 | 不良事件处置费用 | 中医干预相关不适、慢病急性加重急诊/住院费用 | 1.4% |
| 33 | 西医配套及其他成本 | 其他直接成本 | 门诊费、注射费、耗材费等 | 1.4% |

### 三、中西医临床产出与中医特色疗效指标（总权重25%）
| 序号 | 二级指标 | 三级指标 | 指标描述 | 权重 |
|------|---------|---------|---------|------|
| 34 | 西医客观生理结局 | 体重变化 | 干预前后体重差值（kg） | 0.8% |
| 35 | 西医客观生理结局 | BMI变化 | 干预前后BMI差值 | 0.8% |
| 36 | 西医客观生理结局 | 血压达标率 | 收缩压<140且舒张压<90人数及占比 | 1.2% |
| 37 | 西医客观生理结局 | 糖代谢指标变化 | 空腹血糖、糖化血红蛋白干预前后差值 | 1.2% |
| 38 | 西医客观生理结局 | 血脂达标率 | 血脂四项达标人数及占比 | 1.2% |
| 39 | 西医客观生理结局 | 慢病新发病例情况 | 高血压、糖尿病、血脂异常新发人数、相对风险降低率 | 1.6% |
| 40 | 西医客观生理结局 | 心脑血管不良事件 | 心肌梗死、脑卒中、严重并发症发生频次 | 1.2% |
| 41 | 通用生命质量产出 | 健康效用值(QALY) | EQ-5D-5L量表测算质量调整生命年 | 1.6% |
| 42 | 通用生命质量产出 | SF-36生存质量总分 | 生理功能、躯体疼痛、社会功能、精神健康维度得分 | 1.4% |
| 43 | 通用生命质量产出 | EQ-5D各维度受损水平 | 行动、自理、日常活动、疼痛不适、焦虑抑郁分级 | 1.0% |
| 44 | 中医特色治未病疗效产出 | 中医体质改善转归率 | 偏颇质转为平和质、偏颇质积分下降人群占比 | 2.2% |
| 45 | 中医特色治未病疗效产出 | 中医证候总积分下降值 | 干预前后证候积分差值 | 2.0% |
| 46 | 中医特色治未病疗效产出 | 单项中医症状缓解率 | 失眠、乏力、畏寒、腹胀、肩颈酸痛等症状改善人数占比 | 1.8% |
| 47 | 中医特色治未病疗效产出 | 既病防变进展延缓率 | 慢病高危人群进展为确诊慢病比例下降幅度 | 2.0% |
| 48 | 中医特色治未病疗效产出 | 瘥后防复复发率 | 慢病康复人群年度急性复发频次、复发下降比例 | 1.8% |
| 49 | 中医特色治未病疗效产出 | 中医养生依从性 | 规律坚持功法、食疗、外治干预人群占比 | 1.6% |
| 50 | 中医特色治未病疗效产出 | 治未病三级预防阶段适配度 | 未病先防/既病防变/瘥后防复匹配干预场景精准度 | 1.6% |

### 四、间接疾病经济负担指标（总权重15%）
| 序号 | 二级指标 | 三级指标 | 指标描述 | 权重 |
|------|---------|---------|---------|------|
| 51 | 个人间接成本 | 误工损失 | 因病误工工作日×日均薪资（元） | 1.5% |
| 52 | 个人间接成本 | 交通往返成本 | 就医/干预往返交通费用总和（元） | 1.0% |
| 53 | 个人间接成本 | 陪护家属误工成本 | 家属陪护误工折算费用（元） | 1.5% |
| 54 | 个人间接成本 | 养生时间机会成本 | 每日治未病干预投入时间折算经济成本（元） | 1.0% |
| 55 | 社会公共成本 | 医保基金支出节约额 | 慢病医保年度支出减少额（元） | 2.0% |
| 56 | 社会公共成本 | 住院人次节约费用 | 住院减少带来的医保预算结余（元） | 1.5% |
| 57 | 社会公共成本 | 早死生产力损失改善 | 干预后慢病早死导致的社会生产力损失减少额度 | 1.5% |
| 58 | 远期成本效益 | 5年并发症治疗费用减少 | 预测5年内慢病并发症治疗费用减少总额（元） | 2.0% |
| 59 | 远期成本效益 | 急诊就诊频次下降节约 | 急诊次数减少带来的费用节约（元） | 2.0% |

### 五、卫生经济学综合评价核心指标（总权重18%）
| 序号 | 二级指标 | 三级指标 | 指标描述 | 权重 |
|------|---------|---------|---------|------|
| 60 | 成本-效果分析 | 单位慢病新发减少成本 | 每减少1例慢病新发所需投入成本（元/例） | 2.0% |
| 61 | 成本-效果分析 | 单位证候积分下降成本 | 每下降1分中医证候积分所需成本（元/分） | 2.0% |
| 62 | 成本-效果分析 | 住院人次减少单位成本 | 每减少1次住院所需投入成本（元/次） | 2.0% |
| 63 | 成本-效用分析 | 增量成本效用比(ICUR) | 每多获得1个QALY所需增量成本（元/QALY） | 2.0% |
| 64 | 成本-效用分析 | 成本效用比(CUR) | 总成本/总QALYs（元/QALY） | 2.0% |
| 65 | 成本-效益分析 | 净效益(NB) | 干预节约总支出 - 治未病干预投入总成本（元） | 2.0% |
| 66 | 成本-效益分析 | 效益成本比(BCR) | 总效益/总成本 | 2.0% |
| 67 | 最小成本分析 | 同等产出成本差值 | 同等健康产出下中医方案与常规方案总成本差值 | 2.0% |
| 68 | 预算影响分析 | 医保年度支出变化 | 区域人群普及治未病干预后医保年度总支出变化幅度 | 1.0% |
| 69 | 预算影响分析 | 人均年度成本变化 | 干预后人均年度医疗费用变化（元） | 1.0% |

### 六、安全性、满意度与资源消耗辅助指标（总权重7%）
| 序号 | 二级指标 | 三级指标 | 指标描述 | 权重 |
|------|---------|---------|---------|------|
| 70 | 安全性评价 | 不良事件发生率 | 中医外治、中药调理轻度不适、严重不良反应人数占比 | 1.0% |
| 71 | 安全性评价 | 不良事件严重程度分级 | 轻度/中度/重度/危及生命 | 1.0% |
| 72 | 服务满意度评价 | 治未病服务总体满意度 | 非常不满意/不满意/一般/满意/非常满意 | 1.0% |
| 73 | 服务满意度评价 | 各项服务满意度评分 | 体质调理、随访、功法指导、医师服务等分项评分(1-10分) | 1.0% |
| 74 | 卫生资源消耗 | 人均年度门诊次数 | 干预期间人均门诊就诊次数 | 1.0% |
| 75 | 卫生资源消耗 | 人均年度住院频次 | 干预期间人均住院次数 | 1.0% |
| 76 | 卫生资源消耗 | 人均急诊就诊人次 | 干预期间人均急诊次数 | 1.0% |

## 输出格式
请严格按JSON格式返回，不要有任何其他文字：

{
  "document_summary": "对文档内容的简要总结（200字以内）",
  "dimensions": {
    "一、人群基线与中医体质基线指标（15%）": {
      "total_weight": 15,
      "indicators": [
        {"seq": 1, "name": "年龄", "weight": 0.6, "score": 85, "reason": "明确报告了入组年龄均值±标准差"},
        ...
      ]
    },
    "二、治未病干预直接成本指标（20%）": {...},
    "三、中西医临床产出与中医特色疗效指标（25%）": {...},
    "四、间接疾病经济负担指标（15%）": {...},
    "五、卫生经济学综合评价核心指标（18%）": {...},
    "六、安全性、满意度与资源消耗辅助指标（7%）": {...}
  },
  "overall_score": 72.5,
  "overall_assessment": "综合评分结论（200字以内）",
  "suggestions": ["改进建议1", "改进建议2", "改进建议3"]
}
"""


def score_document(content):
    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)

    response = client.chat.completions.create(
        model=DEEPSEEK_MODEL,
        messages=[
            {"role": "system", "content": EVALUATION_CRITERIA},
            {
                "role": "user",
                "content": f"请对以下研究文档进行评分：\n\n{content[:30000]}",
            },
        ],
        temperature=0.1,
        max_tokens=8192,
    )

    result_text = response.choices[0].message.content.strip()

    # Handle possible markdown code fence
    if result_text.startswith("```"):
        lines = result_text.split("\n")
        result_text = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])

    return json.loads(result_text)


@app.route("/")
def index():
    return send_from_directory("templates", "index.html")


@app.route("/api/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "请上传文件"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "请选择文件"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "仅支持 .docx 和 .pdf 格式"}), 400

    original_filename = file.filename
    ext = original_filename.rsplit(".", 1)[1].lower()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = secure_filename(original_filename)
    if not safe_name or "." not in safe_name:
        safe_name = f"upload_{timestamp}.{ext}"
    saved_name = f"{timestamp}_{safe_name}"
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], saved_name)
    file.save(filepath)

    # Extract text
    try:
        content = extract_text(filepath, ext)
        if not content.strip():
            os.remove(filepath)
            return jsonify({"error": "无法提取文档内容，请检查文件格式"}), 400
    except Exception as e:
        os.remove(filepath)
        return jsonify({"error": f"文档解析失败：{str(e)}"}), 400

    # Read binary for DB
    with open(filepath, "rb") as f:
        file_data = f.read()

    # Score via AI
    try:
        score_result = score_document(content)
    except Exception as e:
        return jsonify({"error": f"AI评分失败：{str(e)}"}), 500

    overall = score_result.get("overall_score", 0)
    dim_scores = {}
    for dim_name, dim_data in score_result.get("dimensions", {}).items():
        ind_scores = {f"seq{i['seq']}": i["score"] for i in dim_data.get("indicators", [])}
        dim_scores[dim_name] = {
            "total_weight": dim_data.get("total_weight", 0),
            "indicator_scores": ind_scores,
        }

    db = get_db()
    db.execute(
        """INSERT INTO evaluations (original_filename, file_type, file_content, file_data,
                                      score_json, total_score, dimension_scores, created_at)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
        (
            original_filename,
            ext,
            content,
            file_data,
            json.dumps(score_result, ensure_ascii=False),
            overall,
            json.dumps(dim_scores, ensure_ascii=False),
            datetime.now().isoformat(),
        ),
    )
    db.commit()
    eval_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]

    return jsonify({
        "id": eval_id,
        "total_score": overall,
        "result": score_result,
    })


@app.route("/api/evaluations", methods=["GET"])
def list_evaluations():
    db = get_db()
    rows = db.execute(
        "SELECT id, original_filename, file_type, total_score, created_at "
        "FROM evaluations ORDER BY created_at DESC LIMIT 50"
    ).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/evaluations/<int:eval_id>", methods=["GET"])
def get_evaluation(eval_id):
    db = get_db()
    row = db.execute(
        "SELECT id, original_filename, file_type, score_json, total_score, "
        "dimension_scores, created_at FROM evaluations WHERE id = ?",
        (eval_id,),
    ).fetchone()

    if not row:
        return jsonify({"error": "记录不存在"}), 404

    r = dict(row)
    r["score_json"] = json.loads(r["score_json"])
    r["dimension_scores"] = json.loads(r["dimension_scores"])
    return jsonify(r)


@app.route("/api/evaluations/<int:eval_id>/download", methods=["GET"])
def download_file(eval_id):
    db = get_db()
    row = db.execute("SELECT file_data, original_filename FROM evaluations WHERE id = ?", (eval_id,)).fetchone()
    if not row:
        return jsonify({"error": "记录不存在"}), 404

    from flask import Response
    from urllib.parse import quote
    encoded = quote(row["original_filename"])
    return Response(
        row["file_data"],
        mimetype="application/octet-stream",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded}",
        },
    )


@app.route("/api/evaluations/<int:eval_id>/report", methods=["GET"])
def download_report(eval_id):
    db = get_db()
    row = db.execute(
        "SELECT original_filename, score_json, total_score FROM evaluations WHERE id = ?",
        (eval_id,),
    ).fetchone()
    if not row:
        return jsonify({"error": "记录不存在"}), 404

    from flask import Response
    from urllib.parse import quote
    import io

    result = json.loads(row["score_json"])

    report = io.BytesIO()
    doc = docx.Document()
    doc.styles["Normal"].font.name = "SimSun"
    doc.styles["Normal"].font.size = docx.shared.Pt(11)

    doc.add_heading("中医治未病卫生经济学综合评价报告", 0)
    doc.add_paragraph(f"文献名称：{row['original_filename']}")
    doc.add_paragraph(f"综合得分：{row['total_score']:.1f} / 100")
    doc.add_paragraph(f"评价时间：{row['created_at'] if 'created_at' in row.keys() else ''}")

    doc.add_heading("一、文档概要", level=1)
    doc.add_paragraph(result.get("document_summary", ""))

    doc.add_heading("二、各维度评分明细", level=1)
    for dim_name, dim_data in result.get("dimensions", {}).items():
        dim_weight = dim_data.get("total_weight", 0)
        indicators = dim_data.get("indicators", [])
        dim_score = sum(i.get("score", 0) * i.get("weight", 0) for i in indicators)
        doc.add_heading(f"{dim_name}（{dim_weight}分）→ 得分 {dim_score:.1f}", level=2)

        table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
        hdr = table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "指标名称"
        hdr[2].text = "得分"
        hdr[3].text = "评分依据"
        for ind in indicators:
            row_cells = table.add_row().cells
            row_cells[0].text = str(ind.get("seq", ""))
            row_cells[1].text = ind.get("name", "")
            row_cells[2].text = f"{ind.get('score', 0)}"
            row_cells[3].text = ind.get("reason", "")

    doc.add_heading("三、综合评价", level=1)
    doc.add_paragraph(result.get("overall_assessment", ""))

    doc.add_heading("四、改进建议", level=1)
    for i, s in enumerate(result.get("suggestions", []), 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.save(report)
    report.seek(0)

    base_name = row["original_filename"].rsplit(".", 1)[0]
    report_name = f"评价报告_{base_name}.docx"
    encoded_name = quote(report_name)

    return Response(
        report.getvalue(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}",
        },
    )


if __name__ == "__main__":
    init_db()
    app.run(host="0.0.0.0", port=6007, debug=False)
