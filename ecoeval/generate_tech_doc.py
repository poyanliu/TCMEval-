"""Generate the technical solution document for the EcoEval system."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()


def set_east_asia(run, font_name):
    """Set the East Asian font for a run."""
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)

# ── Global styles ──────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

for level, size, color in [(1, 16, "2E74B5"), (2, 13, "2E74B5"), (3, 11, "404040")]:
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor.from_string(color)
    h._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("1F4E79")
    set_east_asia(run, "黑体")
    return p


def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("595959")
    return p


def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    set_east_asia(run, "宋体")
    # light gray shading
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    p._element.get_or_add_pPr().append(shd)
    return p


# ── Cover ──────────────────────────────────────────────────────
title("中医治未病卫生经济学综合评价系统")
title("技术方案说明书")
subtitle("EcoEval — 技术栈 / 技术路线")

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("文档版本：V1.0\n编制日期：2026-08-13\n部署地址：https://tcmeval.cn/ecoeval/")

doc.add_page_break()

# ── 1. 项目概述 ────────────────────────────────────────────────
doc.add_heading("1. 项目概述", level=1)

doc.add_heading("1.1 项目背景", level=2)
doc.add_paragraph(
    "本系统面向中医治未病卫生经济学综合评价场景，构建一套自动化、智能化的文献/研究报告评价工具。"
    "用户上传研究文档（Word/PDF），系统基于大模型智能解析文档内容，"
    "并依据《中医治未病卫生经济学综合评价指标体系》中的 6 大类 77 项三级指标进行逐项打分，"
    "最终输出加权综合得分、分维度得分明细、综合评价结论与改进建议。"
)

doc.add_heading("1.2 建设目标", level=2)
for goal in [
    "零门槛使用：用户仅需上传文档即可获得结构化评分，无需理解复杂评价规则。",
    "评价标准化：评分规则固化于系统提示词中，保证不同文档评价口径一致。",
    "结果可追溯：每次评分结果与原始文件一并入库，支持历史查询与报告导出。",
    "低运维成本：采用轻量级技术栈，单机即可部署，借助 Cloudflare Tunnel 零成本对外发布。",
]:
    doc.add_paragraph(goal, style="List Bullet")

# ── 2. 技术栈选型 ──────────────────────────────────────────────
doc.add_heading("2. 技术栈选型", level=1)

doc.add_paragraph("系统整体采用前后端分离的轻量级单体架构，核心选型如下：")

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "层级"
hdr[1].text = "技术选型"
hdr[2].text = "选型理由"

tech_rows = [
    ("前端", "原生 HTML + CSS + JavaScript", "无构建工具、无框架依赖，单文件交付，部署简单，加载快"),
    ("后端", "Python 3.12 + Flask 3.1", "轻量、易扩展，与数据科学生态无缝衔接"),
    ("AI 引擎", "DeepSeek (deepseek-chat，V4)", "OpenAI 兼容协议，中文能力强，成本低"),
    ("文档解析", "python-docx + PyMuPDF", "覆盖 Word/PDF 两种主流格式，提取准确"),
    ("数据库", "SQLite (sqlite3)", "零配置嵌入式数据库，单机场景足够，备份简单"),
    ("反向代理", "Nginx 1.18", "统一入口、路径重写、大文件缓冲、超时控制"),
    ("内网穿透", "Cloudflare Tunnel", "无需公网 IP/备案，HTTPS 自动，免费"),
    ("AI SDK", "openai 2.37（兼容客户端）", "以 OpenAI 协议调用 DeepSeek，标准稳定"),
]
for row_data in tech_rows:
    cells = table.add_row().cells
    for i, val in enumerate(row_data):
        cells[i].text = val

doc.add_paragraph()

# ── 3. 系统架构 ────────────────────────────────────────────────
doc.add_heading("3. 系统架构", level=1)

doc.add_heading("3.1 架构图（逻辑）", level=2)
code_block(
    "浏览器 (tcmeval.cn/ecoeval/)\n"
    "        │  HTTPS\n"
    "        ▼\n"
    "Cloudflare Tunnel (cloudflared)\n"
    "        │  /ecoeval/* → localhost:8001\n"
    "        ▼\n"
    "Nginx 反向代理 (0.0.0.0:8001)\n"
    "        │  剥离 /ecoeval 前缀\n"
    "        ▼\n"
    "Flask 应用 (0.0.0.0:6007)\n"
    "        │\n"
    "        ├── 文档解析 (python-docx / PyMuPDF)\n"
    "        ├── AI 评分 (DeepSeek API)\n"
    "        └── SQLite 持久化 (evaluations.db)\n"
)

doc.add_heading("3.2 请求链路", level=2)
doc.add_paragraph(
    "用户访问 https://tcmeval.cn/ecoeval/ 时，请求经 Cloudflare 边缘节点转发至服务器上的 cloudflared 进程；"
    "cloudflared 依据 ingress 规则将 /ecoeval/* 路径转发到本地 Nginx(8001)；"
    "Nginx 通过 location 匹配剥离 /ecoeval 前缀后，将请求代理到 Flask(6007) 的根路径。"
    "这种「隧道 + 反向代理 + 应用」三层结构实现了：路径隔离、前缀归一化、以及大文件与长请求的可靠转发。"
)

# ── 4. 核心模块设计 ────────────────────────────────────────────
doc.add_heading("4. 核心模块设计", level=1)

doc.add_heading("4.1 后端模块（app.py）", level=2)
doc.add_paragraph("后端为单文件 Flask 应用，按功能划分为以下几个部分：")

module_rows = [
    ("文件上传模块", "/api/upload", "接收 multipart 上传，校验扩展名与大小（≤32MB），落盘后解析文本"),
    ("文档解析模块", "extract_text()", "按扩展名分发：.docx 用 python-docx（含段落+表格）；.pdf 用 PyMuPDF 逐页提取"),
    ("AI 评分模块", "score_document()", "构造系统提示词（含 77 项指标全表），调用 DeepSeek 返回 JSON 评分"),
    ("评分存储模块", "SQLite", "评分 JSON、总分、分维度得分、原始文件二进制一并入库"),
    ("历史查询模块", "/api/evaluations", "按时间倒序返回最近 50 条记录摘要"),
    ("详情/下载模块", "/api/evaluations/<id>[/download|/report]", "详情返回完整评分 JSON；download 返回源文件；report 动态生成 Word 报告"),
]
table2 = doc.add_table(rows=1, cols=3)
table2.style = "Light Grid Accent 1"
hdr = table2.rows[0].cells
hdr[0].text = "模块"
hdr[1].text = "路由/函数"
hdr[2].text = "职责"
for row_data in module_rows:
    cells = table2.add_row().cells
    for i, val in enumerate(row_data):
        cells[i].text = val

doc.add_paragraph()

doc.add_heading("4.2 AI 评分流程", level=2)
code_block(
    "1. 上传文档 → 提取纯文本（截断至 30000 字符）\n"
    "2. 组装系统提示词：\n"
    "   - 角色设定：中医治未病卫生经济学评价专家\n"
    "   - 评分规则：每指标 0-100 分，×权重系数=加权得分，总分=Σ加权\n"
    "   - 指标全表：6 大类 77 项（名称/描述/权重）\n"
    "3. 调用 DeepSeek：temperature=0.1（低随机性保证稳定），max_tokens=8192\n"
    "4. 解析返回 JSON：\n"
    "   { document_summary, dimensions{...}, overall_score,\n"
    "     overall_assessment, suggestions[] }\n"
    "5. 总分写入数据库，分维度得分二次汇总存储\n"
)

doc.add_paragraph(
    "为保证输出可解析，提示词强制要求「严格按 JSON 格式返回，不要有任何其他文字」，"
    "并额外处理模型可能包裹的 markdown 代码块（```json ... ```）。"
    "temperature 设为 0.1 以降低评分随机性，保证同一文档多次评分结果基本一致。"
)

doc.add_heading("4.3 前端模块（templates/index.html）", level=2)
doc.add_paragraph("前端为单页应用（SPA），包含以下交互能力：")
for item in [
    "拖拽/点击上传：支持 .docx / .pdf，前端校验格式与大小。",
    "评分进度展示：上传后显示加载动画，评分完成后渲染结果。",
    "结果可视化：环形得分图（conic-gradient 实现）、分维度卡片、指标条形图、改进建议列表。",
    "历史记录：表格展示历史评分，支持「查看详情 / 下载源文件 / 下载报告」三类操作。",
    "响应式布局：CSS Grid 自适应，移动端单列、桌面端多列。",
]:
    doc.add_paragraph(item, style="List Bullet")

# ── 5. 数据库设计 ──────────────────────────────────────────────
doc.add_heading("5. 数据库设计", level=1)

doc.add_paragraph("采用 SQLite 单表存储，表结构如下：")

code_block(
    "CREATE TABLE evaluations (\n"
    "    id INTEGER PRIMARY KEY AUTOINCREMENT,\n"
    "    original_filename TEXT NOT NULL,   -- 原始文件名（含中文）\n"
    "    file_type TEXT NOT NULL,           -- 扩展名 docx/pdf\n"
    "    file_content TEXT NOT NULL,        -- 提取的纯文本\n"
    "    file_data BLOB,                    -- 原始文件二进制\n"
    "    score_json TEXT NOT NULL,          -- 完整评分结果 JSON\n"
    "    total_score REAL NOT NULL,         -- 综合总分\n"
    "    dimension_scores TEXT NOT NULL,    -- 分维度得分 JSON\n"
    "    created_at TEXT NOT NULL           -- 评分时间（ISO8601）\n"
    ");\n"
)

doc.add_paragraph(
    "设计要点：原始文件以 BLOB 形式与评分结果同表存储，保证「结果与文件」强一致、"
    "可追溯、可回迁；score_json 保留 AI 原始完整输出，dimension_scores 为便于快速检索的二次汇总；"
    "未采用外键与多表拆分，符合单机轻量场景，避免过度设计。"
)

# ── 6. 部署方案 ────────────────────────────────────────────────
doc.add_heading("6. 部署方案", level=1)

doc.add_heading("6.1 服务清单", level=2)
deploy_rows = [
    ("Flask 应用", "6007", "nohup python app.py", "核心业务"),
    ("Nginx 反代", "8001", "systemd / nginx", "路径重写 + 大文件缓冲"),
    ("Cloudflare Tunnel", "-", "nohup cloudflared tunnel run <id>", "公网 HTTPS 入口"),
]
table3 = doc.add_table(rows=1, cols=4)
table3.style = "Light Grid Accent 1"
hdr = table3.rows[0].cells
for i, t in enumerate(["服务", "端口", "启动方式", "职责"]):
    hdr[i].text = t
for row_data in deploy_rows:
    cells = table3.add_row().cells
    for i, val in enumerate(row_data):
        cells[i].text = val

doc.add_paragraph()

doc.add_heading("6.2 关键配置", level=2)
doc.add_paragraph("Nginx 反向代理核心配置：")
code_block(
    "location /ecoeval/ {\n"
    "    proxy_pass http://127.0.0.1:6007/;   # 末尾斜杠=剥离前缀\n"
    "    proxy_read_timeout 300s;              # 覆盖 AI 长响应\n"
    "    client_max_body_size 40m;             # 大文件上传\n"
    "}\n"
)

doc.add_paragraph("Cloudflare Tunnel ingress 规则（按顺序匹配，先命中先生效）：")
code_block(
    "ingress:\n"
    "  - hostname: tcmeval.cn\n"
    "    path: /ecoeval/*\n"
    "    service: http://localhost:8001   # 新评价系统\n"
    "  - hostname: tcmeval.cn\n"
    "    path: /api/*\n"
    "    service: http://localhost:8000   # 原有 API\n"
    "  - hostname: tcmeval.cn\n"
    "    service: http://localhost:8501   # Streamlit 主应用\n"
    "  - service: http_status:404\n"
)

doc.add_heading("6.3 启动步骤", level=2)
code_block(
    "# 1. 启动 Flask 应用\n"
    "cd /root/EcoEval && nohup python app.py > /tmp/ecoeval.log 2>&1 &\n"
    "\n"
    "# 2. 启动 Nginx（已配置好 /ecoeval/ 反代）\n"
    "nginx\n"
    "\n"
    "# 3. 启动/重启 Cloudflare Tunnel（应用 ingress 规则）\n"
    "nohup cloudflared tunnel run 9d405ab8-... > /tmp/cloudflared.log 2>&1 &\n"
)

# ── 7. 关键技术难点与解决方案 ──────────────────────────────────
doc.add_heading("7. 关键技术难点与解决方案", level=1)

difficulty_rows = [
    ("中文文件名导致扩展名丢失",
     "secure_filename() 会剥离中文，使「xx.pdf」扩展名丢失引发 IndexError",
     "从原始 file.filename 提取扩展名，再对安全名做兜底"),
    ("下载 502 Bad Gateway",
     "响应头 Content-Disposition 含中文，Werkzeug 用 latin-1 编码失败导致上游连接中断",
     "文件名用 RFC 5987 (filename*=UTF-8''...) 做 URL 编码"),
    ("AI 输出非严格 JSON",
     "模型偶发包裹 markdown 代码块或附加说明文字",
     "剥离代码块围栏后 json.loads 解析，temperature=0.1 降低随机性"),
    ("大文档超出上下文",
     "完整文档可能超出模型上下文窗口",
     "截断至 30000 字符，聚焦关键指标相关内容"),
    ("路径前缀不一致",
     "公网 /ecoeval/ 前缀与 Flask 根路径不匹配",
     "Nginx proxy_pass 末尾斜杠自动剥离前缀，实现路径归一化"),
]
table4 = doc.add_table(rows=1, cols=3)
table4.style = "Light Grid Accent 1"
hdr = table4.rows[0].cells
for i, t in enumerate(["难点", "根因", "解决方案"]):
    hdr[i].text = t
for row_data in difficulty_rows:
    cells = table4.add_row().cells
    for i, val in enumerate(row_data):
        cells[i].text = val

doc.add_paragraph()

# ── 8. 安全与性能 ──────────────────────────────────────────────
doc.add_heading("8. 安全与性能", level=1)

doc.add_heading("8.1 安全措施", level=2)
for item in [
    "上传校验：白名单扩展名（docx/pdf/doc）+ 32MB 大小限制，防止恶意文件。",
    "路径安全：使用 secure_filename 清洗文件名，防止路径穿越。",
    "API Key 保护：DeepSeek 密钥仅存于服务端环境，前端不暴露。",
    "HTTPS 传输：由 Cloudflare 自动提供 TLS，端到端加密。",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("8.2 性能优化", level=2)
for item in [
    "AI 调用采用低随机性 + 限长文本，控制单次响应时间。",
    "SQLite 按需连接（g 对象），请求结束即关闭，避免连接泄漏。",
    "历史查询仅返回摘要字段（不含 BLOB），降低数据库 IO。",
    "前端无重框架，首屏加载快；图表为纯 CSS/原生实现，无额外依赖。",
]:
    doc.add_paragraph(item, style="List Bullet")

# ── 9. 未来扩展方向 ────────────────────────────────────────────
doc.add_heading("9. 未来扩展方向", level=1)
for item in [
    "生产化部署：由 Flask 开发服务器迁移至 Gunicorn/uWSGI 多进程 WSGI。",
    "数据库升级：SQLite 迁移至 PostgreSQL，支持并发写入与多用户隔离。",
    "用户体系：增加账号登录、权限分级、团队协作。",
    "模型可插拔：抽象 LLM 客户端，支持切换多家模型供应商。",
    "批量评价：支持多文件并发评分与横向对比。",
    "评价体系可视化编辑：指标/权重配置化，无需改代码即可调整评价标准。",
]:
    doc.add_paragraph(item, style="List Bullet")

# ── 结语 ────────────────────────────────────────────────────────
doc.add_heading("10. 结语", level=1)
doc.add_paragraph(
    "本系统以「轻量、标准、可追溯」为设计原则，用最小技术栈实现了「文档上传 → 智能解析 → "
    "标准化评分 → 结果持久化 → 报告导出」的完整闭环。"
    "借助 DeepSeek 大模型的语义理解能力，将原本依赖人工经验的中医治未病卫生经济学评价过程自动化，"
    "为中医卫生经济学研究与政策评估提供了高效、一致的工具支撑。"
)

out_path = "/root/EcoEval/EcoEval技术方案说明书.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
